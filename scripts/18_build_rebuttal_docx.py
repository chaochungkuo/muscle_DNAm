#!/usr/bin/env python3
"""Build a clean round-2 rebuttal draft with embedded multi-panel figures.

The document is intentionally generated as a clean Word file rather than a
tracked-changes document. It uses Juliane's 2026-07-22 rebuttal draft as the
scientific base, but reorganizes the response so each reviewer concern is
paired with a complete, defensible answer and a concrete manuscript action.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ANALYSIS_REPO = Path(__file__).resolve().parents[1]
PROJECT_ROOT = ANALYSIS_REPO.parents[1]
OUT_DIR = PROJECT_ROOT / "manuscripts" / "reviewer_round_2" / "to_Juliane_2026-07-23"
FIG_DIR = OUT_DIR / "figures_response"
OUT_DOCX = OUT_DIR / "Review_round_2_rebuttal_clean_multipanel_draft.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color: str = "D9D9D9") -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def style_doc(doc: Document) -> None:
    sec = doc.sections[0]
    sec.top_margin = Inches(0.7)
    sec.bottom_margin = Inches(0.7)
    sec.left_margin = Inches(0.7)
    sec.right_margin = Inches(0.7)
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    styles["Normal"].font.size = Pt(10)
    for style_name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        styles[style_name].font.name = "Arial"
        styles[style_name]._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")


def add_title(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Review round 2 – clean rebuttal draft")
    run.bold = True
    run.font.size = Pt(16)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "Working version for Juliane; based on 260722-Review round 2.docx and updated with consolidated response figures."
    )
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(90, 90, 90)


def add_overall_strategy(doc: Document) -> None:
    add_box(
        doc,
        "Overall response strategy",
        """
The revision follows one consistent logic. First, we acknowledge that this is a retrospective pilot cohort in which disease group, sample source, Sentrix ID, biopsy site, age, sex and tissue-composition features partly overlap. Second, we show the reviewer-requested robustness analyses: metadata-colored PCA/t-SNE, t-SNE parameter sensitivity, expanded PCA, full-matrix label-free correlation heatmap, recomputed subset analyses, covariate sensitivity where estimable, metadata-only classification and leakage-controlled patient-aware supervised learning. Third, we revise the manuscript claims accordingly: the data support disease-group-associated methylation structure in routine diagnostic muscle biopsies, but not a validated clinical diagnostic classifier or a disease-intrinsic methylation signature. ALS versus non-ALS NMA is retained as exploratory and hypothesis-generating.
""",
        "EEF4FF",
    )


def add_box(doc: Document, label: str, text: str, fill: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell)
    p = cell.paragraphs[0]
    r = p.add_run(label)
    r.bold = True
    r.font.size = Pt(9)
    p.add_run("\n")
    for line_i, line in enumerate(text.strip().split("\n")):
        if line_i:
            p.add_run("\n")
        p.add_run(line.strip())
    doc.add_paragraph()


def add_reviewer_comment(doc: Document, text: str) -> None:
    add_box(doc, "Reviewer comment", text, "F7F7F7")


def add_response(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run("Draft response: ")
    r.bold = True
    p.add_run(text.strip())


def add_action(doc: Document, text: str) -> None:
    add_box(doc, "Suggested manuscript action", text, "FFF8E5")


def add_figure(doc: Document, filename: str, caption: str, width: float = 6.5) -> None:
    fig_path = FIG_DIR / filename
    if not fig_path.exists():
        raise FileNotFoundError(fig_path)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(fig_path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.italic = True
    r.font.size = Pt(8)


def add_figure_reference(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run("Figure reference: ")
    r.bold = True
    p.add_run(text.strip())


def add_section(
    doc: Document,
    heading: str,
    reviewer_text: str,
    response_text: str,
    action_text: str,
    figures: list[tuple[str, str]],
    figure_references: list[str] | None = None,
) -> None:
    doc.add_heading(heading, level=2)
    add_reviewer_comment(doc, reviewer_text)
    add_response(doc, response_text)
    add_action(doc, action_text)
    for filename, caption in figures:
        add_figure(doc, filename, caption)
    for reference in figure_references or []:
        add_figure_reference(doc, reference)


def build_doc() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style_doc(doc)
    add_title(doc)
    add_overall_strategy(doc)

    doc.add_heading("Reviewer 2", level=1)

    add_section(
        doc,
        "Reviewer 2.1. Inclusion/exclusion criteria and analytical sample set",
        """
The reviewer states that the inclusion and exclusion criteria are still not sufficiently detailed. Specifically, the reviewer asks how the small non-IBM IIM, NOS and non-ALS NMA groups were selected, whether all candidate cases in the eight-year period were reviewed, whether slides or reports were reviewed before inclusion, how differences among diagnosing pathologists were controlled for, and whether clinical information was evaluated.
""",
        """
We thank the reviewer for requesting clarification. We revised the manuscript to state what is supported by the analytical record. The methylation analyses included all samples that were provided for this study, fulfilled the stated pre-analytical sample requirements, and passed methylation-array quality control. Samples were not selected or excluded on the basis of methylation patterns, clustering, supervised-learning performance, or any expected analytical result. Case groups were defined using the clinicopathological disease-group labels available for the study, and the revised manuscript now avoids implying that the analyzed cohort represents all biopsies or all potential cases from a historical source population.
""",
        """
Methods: state that all provided samples meeting the stated tissue/clinical-information requirements and passing methylation-array QC were analyzed. Do not claim historical candidate-case totals that are not available from the analysis record. Results: describe the final analyzed sample set rather than implying post hoc analytical sample selection.
""",
        [],
    )

    add_section(
        doc,
        "Reviewer 2.2. Cell composition and deconvolution limitation",
        """
The reviewer agrees that longitudinal disease analyses are beyond the scope of the manuscript, but clarifies that the core concern is tissue and cell-type composition. The reviewer asks either for deconvolution estimates to be incorporated into differential methylation models, or for the absence of such analyses to be explicitly discussed as a major limitation.
""",
        """
We agree that bulk muscle-biopsy methylation profiles can reflect variable tissue and cellular composition, including inflammatory infiltrates, fibrosis, necrosis/regeneration and other biopsy-composition effects. We therefore expanded the manuscript limitation to state that the current bulk-tissue data cannot distinguish primary disease-intrinsic methylation changes from secondary changes related to tissue composition. We displayed the available lymphomonocytic infiltration category only as supplied in the metadata table, without independently interpreting or rescoring it. Validated fibrosis, necrosis, fiber-type and denervation scores were not available for all samples and were therefore not inferred or incorporated into differential methylation models.
""",
        """
Discussion/limitations: explicitly state that absence of complete validated tissue-composition and pathology-severity covariates is a major limitation. Methods/results: do not infer unavailable fibrosis, necrosis, fiber-type or denervation scores. Keep pathway interpretation exploratory and linked to bulk-tissue limitations.
""",
        [],
    )

    add_section(
        doc,
        "Reviewer 2.3. ALS versus non-ALS NMA language",
        """
The reviewer states that retaining the ALS versus non-ALS comparison is acceptable, but asks that the relevant language be further softened to avoid overinterpretation given the small number of ALS cases.
""",
        """
We agree. We retained the ALS versus non-ALS NMA comparison, but revised the manuscript to describe this analysis as exploratory and hypothesis-generating. The revised text now states that ALS and non-ALS NMA remain partially overlapping in unsupervised analyses, that the ALS versus non-ALS NMA subset was the least stable disease-pair comparison, and that the comparison is sensitive to individual samples. We removed language implying validated diagnostic separation or robust disease-specific classification for ALS versus non-ALS NMA.
""",
        """
Abstract, Results and Discussion: describe ALS versus non-ALS NMA as exploratory and hypothesis-generating. Do not present ALS/NMA separation as a validated diagnostic finding. Refer to Supplementary Response Figure S5 for the ALS/NMA robustness checks.
""",
        [],
        [
            "Supplementary Response Figure S5 summarizes exploratory ALS versus non-ALS NMA subset PCA/t-SNE and individual-sample influence analyses.",
        ],
    )

    add_section(
        doc,
        "Reviewer 2.4. Proofreading, disease-group terminology and supplementary table reference",
        """
The reviewer states that the revised manuscript contains grammatical and typographic errors and incomplete revisions. The reviewer specifically notes that supervised learning should not be described as predicting a correct diagnosis; the correct term is disease group. The reviewer also notes that the abstract overstates prediction beyond disease group, because non-ALS and non-IBM IIM, NOS are not diagnoses and other differential diagnoses have not been evaluated. Finally, the reviewer notes an incorrect reference to supplementary table xy instead of supplementary tables 2 and 3, and asks for careful proofreading before resubmission.
""",
        """
We agree and revised the wording throughout. Supervised learning is now described as disease-group classification within this selected pilot cohort, not as prediction of a correct diagnosis. The abstract, Results, Figure 6 legend and Discussion now state that classifier performance is internal to the selected cohort and requires validation in independent cohorts. We corrected the supplementary table reference to supplementary tables 2 and 3 and proofread the revised manuscript to remove diagnostic overstatements and inconsistent disease-specific wording.
""",
        """
Global manuscript pass: use disease-group terminology for supervised learning labels; avoid diagnostic-classifier validation language; correct the supplementary table reference; remove inconsistent overstatements and typographic errors.
""",
        [],
    )

    doc.add_heading("Reviewer 1", level=1)
    add_section(
        doc,
        "1. Overall concern: strong clustering may reflect confounding",
        """
Reviewer 1 states that the cluster separation in t-SNE and PCA is unexpectedly strong for a small, heterogeneous cohort and asks the authors to establish that the structure is robustly associated with disease groups rather than technical, demographic, or tissue-composition variables. The reviewer specifically highlights age, sex, biopsy site, comorbidities, disease duration, tissue handling, batch/Sentrix, sample source, and bulk-tissue cellular composition.
""",
        """
We thank the reviewer for raising these important points. We agree that bulk DNA methylation profiles from muscle biopsies are influenced by tissue composition, including variable contributions from muscle fibers, adipose tissue, connective tissue, inflammatory infiltrates, necrotic or regenerating fibers, and fibrosis. We also agree that sample source, Sentrix array, biopsy site, age and sex partly overlap with disease group in this retrospective cohort. We therefore revised the interpretation throughout the manuscript. The additional analyses support disease-group-associated methylation structure, while showing that this structure cannot be interpreted as a purely disease-intrinsic or cell-type-specific methylation signature. Instead, the observed profiles likely reflect the composite methylation phenotype of routine diagnostic muscle biopsies. We now present this work as exploratory and hypothesis-generating, requiring validation in larger, prospectively collected and independently processed cohorts.
""",
        """
In Results and Discussion, avoid wording that implies disease-intrinsic specificity. Replace claims such as "diagnostic methylation signature" with "disease-group-associated methylation structure" or "composite methylation profile of diagnostic muscle biopsies". Add a limitation stating that disease group, sample source, Sentrix ID, biopsy site and tissue-composition variables are structurally linked in this retrospective pilot cohort and cannot be fully disentangled.
""",
        [
            (
                "Response_Figure_1_unsupervised_structure.png",
                "Response Figure 1. Unsupervised methylation structure and cohort variables. Baseline PCA, t-SNE, PC-metadata associations and full-matrix correlation heatmap show group-associated structure while documenting overlap with cohort variables.",
            )
        ],
    )

    add_section(
        doc,
        "2. t-SNE default PCA, direct pca=FALSE, and parameter sensitivity",
        """
The reviewer notes that Rtsne defaults include pca=TRUE and initial_dims=50, so the baseline t-SNE is calculated from the first 50 internally derived principal components rather than directly from all probes. The reviewer asks for this to be stated explicitly and for sensitivity analyses using direct full-matrix t-SNE, different initial dimensions, perplexities and random seeds.
""",
        """
We thank the reviewer for pointing out the Rtsne default. We now state explicitly that the baseline t-SNE used pca=TRUE and initial_dims=50, meaning that Rtsne first reduced the input matrix to 50 centered, unscaled principal components before computing the two-dimensional t-SNE embedding. To test whether the visualization depended on this default or on a single arbitrary parameter setting, we repeated t-SNE across multiple initial dimensions, perplexities and random seeds, and also performed a direct full-matrix t-SNE with pca=FALSE. These sensitivity analyses show that the major separation pattern is reproducible across a range of t-SNE settings and is therefore not explained by a single random seed or one arbitrary parameter choice. However, the exact position and distance between clusters vary between runs, as expected for t-SNE. We therefore use t-SNE as a qualitative visualization of group-associated structure, not as a quantitative measure of global distances between disease groups.
""",
        """
Methods: explicitly state Rtsne parameters, including pca=TRUE and initial_dims=50 for the baseline analysis, and state that no label-informed probe selection was applied before PCA or t-SNE. Results: add one sentence that sensitivity analyses reproduced the major separation pattern across parameter settings, while exact t-SNE geometry varied. Discussion: avoid quantitative interpretation of t-SNE distances.
""",
        [
            (
                "Response_Figure_2_tsne_pca_sensitivity.png",
                "Response Figure 2. t-SNE and PCA sensitivity analyses. The major structure persists across t-SNE settings and PCA scaling choices, but t-SNE geometry remains parameter-dependent and should be interpreted qualitatively.",
            )
        ],
    )

    add_section(
        doc,
        "3. Metadata and confounder coloring",
        """
The reviewer asks that the primary t-SNE, t-SNE sensitivity analyses and PCA coordinates be displayed with samples colored separately by potential confounder variables, including dataset source, Sentrix/chip, age, sex, biopsy site, inflammatory-cell estimates, fiber-type estimates, denervation estimates and pathology severity where available.
""",
        """
As suggested, we now provide PCA and t-SNE plots colored by available metadata variables, including disease group, sample source, Sentrix ID, age, sex, biopsy-site group and lymphomonocytic infiltration category where provided. The lymphomonocytic infiltration category is displayed only as supplied in the metadata table; we do not independently interpret the scoring scheme or convert uncertain annotations into low/intermediate/high categories ourselves. Requested variables not available in the provided metadata table, including validated fibrosis, necrosis, fiber-type and denervation scores, were not inferred or modeled. These visualizations show that several cohort variables overlap with the disease-group structure, which is now explicitly reflected in the revised interpretation.
""",
        """
Add supplementary response figures showing PCA and t-SNE colored by available metadata. In the rebuttal and manuscript, clearly distinguish available metadata from unavailable variables. Display the lymphomonocytic infiltration category only as supplied, without independently interpreting its pathology meaning. Do not invent low/intermediate/high scores for lymphomonocytes, fibrosis, necrosis, fiber type, denervation or other histological features unless Juliane confirms the scoring.
""",
        [],
        [
            "Supplementary Response Figure S1 shows PCA colored by each available metadata variable.",
            "Supplementary Response Figure S2 shows t-SNE colored by each available metadata variable.",
        ],
    )

    add_section(
        doc,
        "4. PCA evaluation, scaling and PC-metadata association",
        """
The reviewer asks for more detailed PCA evaluation, including scree and cumulative-variance plots, multiple PC projections, evaluation beyond PC1-PC2, and comparison with unscaled PCA because scale.=TRUE gives all CpGs equal variance weight. The reviewer also asks for quantitative association of leading PCs with diagnosis and potential confounders.
""",
        """
We expanded the PCA analysis in four ways. First, we added scree and cumulative-variance plots to show how much methylation variance is captured by the leading PCs. Second, we plotted PC1-PC2, PC1-PC3 and PC2-PC3 to test whether the observed structure is restricted to the first two PCs or also appears in other major components. Third, we annotated the strongest positive and negative CpG loadings of the leading PCs to support interpretation of whether major components reflected disease group, cohort variables or technical structure. Fourth, we repeated PCA using centered but unscaled M-values to test whether the observed structure depends on unit-variance scaling of CpGs. These analyses show that disease-group-associated structure is present in PCA, but that leading PCs also overlap with cohort variables such as source, Sentrix ID, age, sex and biopsy site. We therefore retained the PCA result but revised the interpretation to avoid presenting the PCs as disease-only axes.
""",
        """
Add Methods detail for scaled PCA (scale.=TRUE) and unscaled PCA sensitivity analysis (scale.=FALSE). Add Results text summarizing scree/cumulative variance, PC-metadata association and annotated leading positive/negative CpG loadings. Replace ambiguous text such as "the unscaled analysis supports the presence of structure" with explicit wording: "the group-associated pattern was also visible in unscaled PCA, but leading PCs remained associated with cohort variables." Do not claim that PCA axes are disease-only axes.
""",
        [],
        [
            "Supplementary Response Figure S3 provides the complete PCA sensitivity analysis, including scree/cumulative variance, multiple PC projections and unscaled PCA.",
        ],
    )

    add_section(
        doc,
        "5. Label-free full-matrix correlation heatmap and replacement of old Figure 1F",
        """
The reviewer states that the previous heatmap was misleading when shown together with t-SNE because probes were selected using diagnostic labels. The reviewer asks for a sample-to-sample correlation heatmap calculated from the full post-QC M-value matrix, with samples ordered by unsupervised hierarchical clustering and diagnostic/confounder annotations added only after clustering.
""",
        """
We agree that the previous top-CpG heatmap was label-informed and should not be presented as independent unsupervised evidence. We replaced it with a sample-to-sample Pearson correlation heatmap calculated from the complete post-QC M-value matrix. Samples were ordered by hierarchical clustering using 1 minus Pearson correlation and complete linkage, and diagnostic and metadata annotations were added only after clustering. The previous label-informed top-CpG heatmap, if retained, will be described only as a supervised descriptive visualization and not as evidence of unsupervised clustering.
""",
        """
Replace Figure 1F with the full-matrix sample-to-sample correlation heatmap. Add a complete annotation legend for all colors. Move or re-label any top-CpG heatmap as supervised/descriptive only, or remove it from the main unsupervised figure if space is limited.
""",
        [],
        [
            "Response Figure 1, panels D–E, show the full-matrix label-free correlation heatmap and annotation legend. The figure is embedded once above to avoid duplicating the same multi-panel figure in the main rebuttal."
        ],
    )

    add_section(
        doc,
        "6. Subset analyses requested by the reviewer",
        """
The reviewer asks for PCA and t-SNE recomputed within subsets rather than by removing points from the full embedding: excluding MMC, excluding controls, in-house samples only, IBM versus non-IBM IIM, and ALS versus non-ALS NMA. The reviewer also asks whether apparent separation is robust to individual samples.
""",
        """
We recomputed PCA and t-SNE within each requested subset: excluding MMC, excluding controls, in-house data, IBM versus non-IBM IIM, and ALS versus non-ALS NMA. These analyses test whether the full-cohort structure is driven by one dominant group or by inclusion of external/public data. We use the term in-house data as an analysis label for the non-public/non-GEO methylation samples generated within this study; we do not use it to imply a single-source or otherwise homogeneous cohort. The subset analyses showed that the full-cohort pattern was stronger than several clinically focused subsets, and ALS versus non-ALS NMA was the least stable comparison. We therefore retained this comparison only as exploratory and hypothesis-generating.
""",
        """
Define "in-house data" as the non-public/non-GEO methylation samples generated within this study. Do not imply that this subset is single-source or otherwise homogeneous. In Results and Discussion, state that ALS versus non-ALS NMA is exploratory and sensitive to individual samples. Do not present subset separation as independent diagnostic proof.
""",
        [],
        [
            "Supplementary Response Figure S4 provides PCA and t-SNE recomputed independently within each requested subset, including the in-house-data subset and clinically focused ALS/non-ALS NMA and IBM/non-IBM IIM comparisons.",
            "Supplementary Response Figure S5 focuses on the exploratory ALS versus non-ALS NMA robustness checks, including subset PCA/t-SNE and individual-sample influence analysis.",
        ],
    )

    add_section(
        doc,
        "7. Downstream analyses after confounding assessment",
        """
The reviewer states that the unsupervised analyses should be resolved before downstream disease-specific claims can be interpreted confidently. The reviewer cautions that PCA/t-SNE findings should not be used as unqualified evidence that the methylation data contain intrinsic diagnostic structure and asks the authors to reassess differential methylation, label-informed heatmaps, classifiers and biological interpretation in light of confounding.
""",
        """
The additional analyses support the presence of disease-group-associated methylation structure, while also showing that cohort variables such as source, Sentrix ID, age, sex and biopsy site partly overlap with disease group in this retrospective pilot cohort. Where statistically possible, we fitted differential methylation models adjusted for age/sex or biopsy site and found that several disease-group-associated results persisted, whereas the ALS versus non-ALS NMA contrast was more sensitive to adjustment. We also rebuilt supervised learning using patient-aware splitting and training-only feature selection, and added metadata-only classification as a confounding diagnostic. These analyses do not negate the disease-group-associated findings; rather, they define their appropriate scope. We retained the differential methylation, gene set enrichment and supervised-learning findings, but revised the manuscript to avoid overclaiming clinical diagnostic performance, disease-intrinsic specificity or final mechanistic proof from pathway enrichment.
""",
        """
Revise downstream Results and Discussion to state that supervised learning supports disease-group-associated signal under leakage-controlled analysis, but is not a validated clinical diagnostic classifier. Describe metadata-only classification as a confounding diagnostic, not as a disease classifier. Keep differential methylation and gene set enrichment interpretation exploratory, use candidate pathway language, and state that pathway findings may reflect both disease-associated biology and tissue-composition differences in bulk biopsy material.
""",
        [
            (
                "Response_Figure_3_downstream_robustness.png",
                "Response Figure 3. Downstream robustness checks, including covariate-adjusted differential methylation where estimable, metadata-only classification and patient-aware supervised learning. ALS/non-ALS NMA subset and influence analyses are provided separately in Supplementary Response Figure S5.",
            )
        ],
    )

    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_heading("Short wording checklist for final manuscript consistency", level=1)
    checks = [
        'Use "disease group" rather than "diagnosis" when referring to machine-learning labels.',
        'Use "in-house data" for the non-public/non-GEO methylation samples generated within this study; avoid wording that implies a homogeneous single-source subset.',
        'Use "disease-group-associated methylation structure" rather than disease-only or disease-intrinsic signature wording unless independently validated.',
        'State that t-SNE is qualitative and parameter-dependent; do not interpret global distances quantitatively.',
        'State that bulk-tissue methylation cannot separate primary disease-intrinsic epigenetic changes from tissue-composition effects.',
        'State that ALS versus non-ALS NMA is exploratory because of small sample size and sample influence.',
    ]
    for check in checks:
        doc.add_paragraph(check, style="List Bullet")

    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    build_doc()

from __future__ import annotations

import importlib.util
import re
import shutil
import subprocess
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Inches, Pt
from PIL import Image


ROOT = Path(__file__).resolve().parents[1]
PROJECT = ROOT.parents[1]
MANUSCRIPTS = PROJECT / "manuscripts"
PACKAGE = MANUSCRIPTS / "reviewer_round_2" / "to_Juliane_2026-07-19"
ZIP = PACKAGE.with_suffix(".zip")
WEB_FIG = ROOT / "figures" / "web"
MAIN_FIG = ROOT / "figures" / "main"
SUPP_FIG = ROOT / "figures" / "supplementary"
FINAL_DRAFTS = MANUSCRIPTS / "reviewer_round_2" / "final_drafts"
SOURCE_REPORT = MANUSCRIPTS / "reviewer_round_2" / "to_Juliane_2026-07-16" / "reviewer_response_workbook.html"


def load_workbook_module():
    path = ROOT / "scripts" / "15_build_reviewer_response_workbook.py"
    spec = importlib.util.spec_from_file_location("reviewer_workbook", path)
    module = importlib.util.module_from_spec(spec)
    assert spec.loader is not None
    spec.loader.exec_module(module)
    return module


def reset_package() -> None:
    if PACKAGE.exists():
        shutil.rmtree(PACKAGE)
    (PACKAGE / "figures").mkdir(parents=True)


def paragraph_with_label(doc: Document, label: str, text: str, highlight=None) -> None:
    p = doc.add_paragraph()
    r = p.add_run(label)
    r.bold = True
    if highlight is not None:
        r.font.highlight_color = highlight
    p.add_run(text)


def add_highlighted_paragraph(doc: Document, text: str, highlight=WD_COLOR_INDEX.YELLOW) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.highlight_color = highlight


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value


def safe_png_from_tiff(source: Path, destination: Path, max_size=(1800, 1300)) -> Path:
    destination.parent.mkdir(parents=True, exist_ok=True)
    with Image.open(source) as img:
        img = img.convert("RGB")
        img.thumbnail(max_size, Image.Resampling.LANCZOS)
        img.save(destination, optimize=True)
    return destination


def figure_png_for_stem(stem: str, temp_dir: Path) -> Path | None:
    web = WEB_FIG / f"{stem}.png"
    if web.exists():
        return web
    source = SUPP_FIG / f"{stem}.tiff"
    if source.exists():
        return safe_png_from_tiff(source, temp_dir / f"{stem}.png")
    source = MAIN_FIG / f"{stem}.tiff"
    if source.exists():
        return safe_png_from_tiff(source, temp_dir / f"{stem}.png")
    return None


def copy_figure_pdf(stem: str) -> None:
    for source_dir in [MAIN_FIG, SUPP_FIG]:
        pdf = source_dir / f"{stem}.pdf"
        if pdf.exists():
            shutil.copy2(pdf, PACKAGE / "figures" / pdf.name)
            return


def stems_for_figure_field(field: str) -> list[str]:
    explicit = re.findall(r"([A-Za-z0-9]+[A-Za-z0-9_]*?)\.pdf", field)
    stems = [Path(x).stem for x in explicit]
    if "pca_scaled_by_*.pdf" in field:
        stems.extend(sorted(p.stem for p in SUPP_FIG.glob("pca_scaled_by_*.pdf")))
    if "tsne_baseline_by_*.pdf" in field:
        stems.extend(sorted(p.stem for p in SUPP_FIG.glob("tsne_baseline_by_*.pdf")))
    if "subset_" in field and "*.pdf" in field:
        stems.extend(sorted(p.stem for p in SUPP_FIG.glob("subset_*_PCA.pdf")))
        stems.extend(sorted(p.stem for p in SUPP_FIG.glob("subset_*_tSNE.pdf")))
    # Preserve order while removing duplicates.
    seen = set()
    unique = []
    for stem in stems:
        if stem not in seen:
            seen.add(stem)
            unique.append(stem)
    return unique


def build_rebuttal_docx(chunks: list[dict]) -> Path:
    out = PACKAGE / "02_rebuttal_letter_draft_with_figures.docx"
    temp_dir = PACKAGE / "_docx_preview_temp"
    temp_dir.mkdir(exist_ok=True)

    doc = Document()
    doc.add_heading("Point-by-point rebuttal letter draft", 0)
    add_highlighted_paragraph(
        doc,
        "Color key: yellow text marks information that Juliane/clinical coauthors must confirm before journal submission.",
    )
    doc.add_paragraph(
        "This draft follows the reviewer-round-2 workbook. Reviewer comments are included in full, followed by draft response text and the figures recommended for insertion in the rebuttal or supplement."
    )

    for chunk in chunks:
        doc.add_page_break()
        doc.add_heading(f"{chunk['reviewer']} — {chunk['title']}", level=1)
        doc.add_heading("Reviewer comment", level=2)
        doc.add_paragraph(chunk["quote"])

        doc.add_heading("Draft response", level=2)
        p = doc.add_paragraph(chunk["draft"])
        if "JULIANE TO INSERT" in chunk["draft"] or "Validated" in chunk["draft"] and "not available" in chunk["draft"]:
            for run in p.runs:
                if "JULIANE" in run.text:
                    run.font.highlight_color = WD_COLOR_INDEX.YELLOW

        figures = chunk.get("rebuttal_figures", [])
        doc.add_heading("Figures to include", level=2)
        if figures:
            rows = [[f["where"], f["figure"], f["caption"]] for f in figures]
            add_table(doc, ["Where", "Figure file(s)", "Caption"], rows)
            for item in figures:
                stems = stems_for_figure_field(item["figure"])
                if not stems:
                    doc.add_paragraph(f"Figure note: {item['figure']}")
                    continue
                for stem in stems:
                    copy_figure_pdf(stem)
                    png = figure_png_for_stem(stem, temp_dir)
                    if png is None:
                        doc.add_paragraph(f"Figure file not found for embedding: {stem}")
                        continue
                    doc.add_paragraph(stem.replace("_", " "))
                    try:
                        doc.add_picture(str(png), width=Inches(5.9))
                    except Exception as exc:
                        doc.add_paragraph(f"Could not embed {stem}: {exc}")
        else:
            doc.add_paragraph("No separate figure is required for this response.")

    doc.save(out)
    shutil.rmtree(temp_dir, ignore_errors=True)
    return out


def highlight_paragraph(paragraph, color=WD_COLOR_INDEX.BRIGHT_GREEN) -> None:
    if not paragraph.runs:
        paragraph.add_run(paragraph.text)
    for run in paragraph.runs:
        run.font.highlight_color = color


def build_highlighted_manuscript() -> Path:
    source = FINAL_DRAFTS / "Bremer_manuscript_reviewer_round_2_clean_draft.docx"
    subprocess.run(["python", "scripts/11_build_revision_documents.py"], cwd=ROOT, check=True)
    doc = Document(source)

    key = doc.paragraphs[0].insert_paragraph_before()
    r = key.add_run(
        "COLOR KEY: green/turquoise highlights mark revised analytical manuscript text; yellow highlights mark Juliane/clinical-author decision items or placeholders."
    )
    r.bold = True
    r.font.highlight_color = WD_COLOR_INDEX.YELLOW
    r.font.size = Pt(11)

    analytical_markers = [
        "Disease group-associated DNA methylation patterns",
        "cross-sectional pilot study",
        "Unsupervised PCA, t-SNE and sample-correlation analyses",
        "larger, prospectively balanced and independently validated cohorts",
        "not as a clinically validated diagnostic test",
        "The original array preprocessing was performed",
        "Baseline t-SNE used Rtsne",
        "For exploratory supervised classification",
        "Metadata-only classifiers were evaluated",
        "Disease-group-associated methylation structure is accompanied",
        "A label-free full-matrix sample-correlation heatmap",
        "Exploratory comparison between related disease groups",
        "The distinction between ALS-associated",
        "Exploratory supervised learning classifies",
        "We evaluated whether methylation data could classify",
        "Figure 6. Patient-aware exploratory classification",
        "In this pilot cohort",
        "The absence of cell-composition-adjusted analyses is therefore a major limitation",
        "ALS and non-ALS NMA were not unequivocally separated",
        "this pilot study demonstrates",
        "The findings require validation",
        "Histological features are included in Supplementary Tables 2 and 3",
    ]
    yellow_markers = [
        "JULIANE DECISION NEEDED",
        "AUTHORS TO CONFIRM",
        "xxx",
        "depending on ongoing analysis",
    ]
    for paragraph in doc.paragraphs:
        text = paragraph.text
        if any(marker in text for marker in yellow_markers):
            highlight_paragraph(paragraph, WD_COLOR_INDEX.YELLOW)
        elif any(marker in text for marker in analytical_markers):
            highlight_paragraph(paragraph, WD_COLOR_INDEX.TURQUOISE)

    out = PACKAGE / "03_manuscript_clean_highlighted_draft.docx"
    doc.save(out)
    return out


def copy_minimal_deliverables() -> None:
    shutil.copy2(SOURCE_REPORT, PACKAGE / "01_reviewer_response_report.html")
    # Include all reviewer-round figure PDFs in one folder. This keeps the top-level
    # package compact while making manuscript/rebuttal figure assembly complete.
    for source_dir in [MAIN_FIG, SUPP_FIG]:
        for pdf in sorted(source_dir.glob("*.pdf")):
            shutil.copy2(pdf, PACKAGE / "figures" / pdf.name)
    readme = """JULIANE COMPACT REVISION PACKAGE
Version: 2026-07-19

Main files:
1. 01_reviewer_response_report.html — main readable report linking reviewer comments, analyses, rebuttal drafts and manuscript modifications.
2. 02_rebuttal_letter_draft_with_figures.docx — point-by-point rebuttal draft with embedded figures where recommended.
3. 03_manuscript_clean_highlighted_draft.docx — clean revised manuscript draft; analytical revisions are highlighted and unresolved Juliane/author decisions are highlighted.

Manuscript base:
- 03_manuscript_clean_highlighted_draft.docx is generated from manuscripts/Bremer_manuscript_changes not marked.docx.

Supporting files:
- figures/ — PDF versions of main and supplementary figures needed for rebuttal, supplement or manuscript figure assembly.

No raw methylation data or large intermediate files are included.
"""
    (PACKAGE / "00_README.txt").write_text(readme, encoding="utf-8")


def zip_package() -> None:
    if ZIP.exists():
        ZIP.unlink()
    with zipfile.ZipFile(ZIP, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=9) as archive:
        for path in sorted(PACKAGE.rglob("*")):
            if path.is_file():
                archive.write(path, PACKAGE.name / path.relative_to(PACKAGE))
    with zipfile.ZipFile(ZIP) as archive:
        bad = archive.testzip()
        if bad is not None:
            raise RuntimeError(f"Corrupt ZIP member: {bad}")


def main() -> None:
    workbook = load_workbook_module()
    reset_package()
    copy_minimal_deliverables()
    build_rebuttal_docx(workbook.CHUNKS)
    build_highlighted_manuscript()
    zip_package()
    print(PACKAGE)
    print(ZIP)


if __name__ == "__main__":
    main()

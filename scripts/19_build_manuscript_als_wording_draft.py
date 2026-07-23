#!/usr/bin/env python3
"""Build a clean highlighted manuscript draft with softened ALS/NMA wording.

The goal is not to comprehensively rewrite the manuscript, but to make the
ALS versus non-ALS NMA and supervised-learning language internally consistent
with the round-2 rebuttal strategy:

- retain ALS/NMA analyses,
- describe them as exploratory/hypothesis-generating,
- avoid diagnostic-validation language,
- use disease-group classification terminology.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.ns import qn
from docx.shared import Pt


ANALYSIS_REPO = Path(__file__).resolve().parents[1]
PROJECT_ROOT = ANALYSIS_REPO.parents[1]
MANUSCRIPTS = PROJECT_ROOT / "manuscripts"
BASE_DOCX = MANUSCRIPTS / "Bremer_manuscript_changes not marked.docx"
OUT_DIR = MANUSCRIPTS / "reviewer_round_2" / "to_Juliane_2026-07-23"
OUT_DOCX = OUT_DIR / "Bremer_manuscript_clean_highlighted_ALS_NMA_wording_draft.docx"
CHANGE_LOG = OUT_DIR / "MANUSCRIPT_ALS_NMA_WORDING_CHANGES.md"


REPLACEMENTS = {
    1: (
        "Title: avoid disease-group-specific overclaim.",
        "Disease group-associated DNA methylation patterns",
    ),
    48: (
        "Abstract: replace overstatement about diagnosis prediction with disease-group classification and validation caveat.",
        "Unsupervised t-SNE analysis and full-matrix sample-to-sample correlation analysis revealed disease-group-associated methylation structure, with partial overlap between ALS and non-ALS NMA. Using leakage-controlled supervised learning, methylation data classified disease groups in held-out samples within this selected pilot cohort, but these internal performance estimates require validation in independent cohorts. Gene set enrichment analysis pointed at commonly dysregulated pathways, including cytoskeletal maintenance, cell adhesion, muscle and neural development, Wnt-signaling as well as proteostasis. In inflammatory myopathies, sites linked to immune system activation were hypomethylated. Analysis of CpG methylation of individual preselected genes involved in ALS and inflammatory myopathies, respectively, pointed at certain genes, whose differential expression might be in part regulated by methylation, including T cell differentiation and cytotoxicity markers. Correlation of CpG site methylation and expression data in IBM revealed a potential mechanism involving epigenetic regulation of DMWD and HDAC4 expression through which muscle stem cells and regeneration can be controlled in IBM. Analyzing RNA expression data, we found further evidence that stem cells fail to activate the regeneration program in IBM.",
    ),
    59: (
        "Introduction aim: avoid claiming disease-specific diagnostic methylome patterns as the established target.",
        "Non-neoplastic muscle disorders are a heterogeneous group of disorders, including genetic, metabolic and toxic diseases as well as autoimmune-mediated inflammatory myopathies and neurogenic atrophy. Although some epigenetic changes have been described for example in dermatomyositis [53, 64], a clear correlation between the muscle methylome and specific muscular disorders has not been established yet. In the present study, we aimed at identifying disease-group-associated methylome patterns in non-neoplastic muscle pathology. Furthermore, we sought to determine whether the methylome alterations are pointing at pathophysiological mechanisms. We also applied leakage-controlled supervised learning to explore whether methylome data can classify disease groups within this selected pilot cohort.",
    ),
    64: (
        "Methods sample source: use in-house diagnostic material wording and avoid implying a homogeneous source subset.",
        "Fresh frozen human skeletal muscle biopsy material was obtained from available in-house diagnostic material: skeletal muscle biopsies from different anatomical sites diagnosed as neurogenic muscular atrophy (ALS NMA and non-ALS NMA) or inflammatory myopathy (inclusion body myopathy (IBM), idiopathic inflammatory myopathy (IIM) without inclusion, not otherwise specified, here termed non-IBM-IIM, NOS) - or from controls from the MALICoT study[18]. We also included previously published data on multi-minicore myopathy [4]. For case selection, we searched available in-house diagnostic cases diagnosed based on neuropathological criteria with IBM, polymyositis/ non-IBM-IIM, NOS, or neurogenic atrophy. IBM cases were defined as inflammatory myopathies characterized by invasion of intact muscle fibers by CD8-positive T cells and with degenerative features, i.e. rimmed vacuoles/inclusions (phospho-TDP43 and p62 positive on immunohistochemistry, confirmed by the detection of abnormal autophagic vacuoles and tubulofilamentous inclusions by electron microscopy). Non-IBM-IIM, NOS cases were inflammatory myopathies with invasion of intact muscle fibers and without rimmed vacuoles/inclusions. Neurogenic atrophy cases were further subdivided based on clinical information regarding the cause (ALS versus non-ALS). The methylation analysis included all samples provided for this study that had sufficiently sized, good-quality frozen muscle tissue, sufficient clinical information and passed methylation-array quality control. No samples were selected or excluded on the basis of methylation patterns, clustering, supervised-learning performance or expected analytical results. Cases with documented prior immunosuppressive treatment were excluded before methylation analysis. Initially, we obtained 26 control samples from the MALICoT study. Histological assessment revealed unspecific histological alterations in two of these cases, which were therefore excluded from this study.",
    ),
    70: (
        "Methods analysis source labels: define in-house data explicitly for reviewer-requested subset analysis.",
        "Array data analysis was performed using R v.4.3.3, using a number of packages from Bioconductor and other repositories. Raw signal intensities were obtained from IDAT files using the minfi R package [2]. Each sample was individually normalised by performing background correction for both colour channels. Subsequently, several filtering criteria were applied to the initial CpG sites (865,859): removal of probes targeting X and Y chromosomes (18,541), removal of probes containing single nucleotide polymorphism within five base pairs spanning and within the targeted CpG site (46,299), and probes with bad quality (23,539). Normalization was performed using the preprocessQuantile function in minfi [62]. After all filtering steps, 771,381 CpG probes remained for downstream analysis. For reviewer-requested source-stratified sensitivity analyses, we used the term in-house data as an analysis label for the non-public/non-GEO methylation samples generated within this study; this label was not intended to imply a single-source or otherwise homogeneous cohort. The M values of the full filtered probe set were used for principal component analysis (PCA) and t-SNE (Rtsne package v0.17). The primary PCA was computed using prcomp on centered and unit-variance-scaled CpGs (scale.=TRUE), so that the analysis assessed broad structure across the filtered probe set rather than being dominated by the highest-variance CpGs. As a sensitivity analysis, PCA was repeated using centered but unscaled M-values (scale.=FALSE), and scree/cumulative-variance plots together with PC1-PC2, PC1-PC3 and PC2-PC3 projections were inspected. Associations between leading PCs and available metadata variables were quantified to evaluate whether major PCs were related to disease group, sample source, Sentrix ID, age, sex or biopsy site. Baseline t-SNE was performed with Rtsne package v0.17 using perplexity = 15, theta = 0.5, dims = 2, set.seed(42), and the Rtsne defaults pca=TRUE and initial_dims=50; thus, Rtsne first reduced the full filtered M-value matrix to 50 internally derived centered, unscaled PCs before computing the two-dimensional embedding. No differential-methylation-based or other label-informed probe selection was applied before PCA or t-SNE. To assess t-SNE robustness, we repeated t-SNE across multiple initial dimensions, perplexities and random seeds, and also performed direct full-matrix t-SNE with pca=FALSE. For Figure 1F, we calculated a label-free sample-to-sample Pearson correlation heatmap from the complete post-QC M-value matrix. Samples were ordered by unsupervised hierarchical clustering using 1 minus Pearson correlation and complete linkage; diagnostic and metadata annotations were added only after clustering. Differential methylation analysis was conducted using the limma R package [51] on the M values of all CpG sites with the annotation on human genome hg19. The adjusted p-values were corrected using the Benjamini-Hochberg procedure to calculate the False Discovery Rate (FDR). Gene Set Enrichment Analysis (GSEA) was performed using the clusterProfiler R package [69]. Customized heatmaps generated from CpG sites selected from limma differential methylation results were treated as label-informed downstream descriptive visualizations and not as independent unsupervised clustering evidence.",
    ),
    71: (
        "Methods heading: replace diagnosis prediction with disease-group classification.",
        "Supervised learning for disease-group classification",
    ),
    85: (
        "Results sample description: use in-house diagnostic cases and avoid source-label ambiguity.",
        "The final in-house diagnostic sample set comprised 41 fresh-frozen muscle biopsies with well-defined neuropathological disease groups of either (1) inflammatory myopathies (IIM, n=22), including inclusion body myositis (IBM, n=16, Fig. 1a) and idiopathic inflammatory myopathy without inclusions, not otherwise specified, here termed non-IBM-IIM, NOS (see above; n=6) or (2) neurogenic atrophy (n= 19, Fig. 1b-d), including six cases of clinically diagnosed amyotrophic lateral sclerosis (ALS) and 13 cases of denervation due to other causes (non-ALS NMA). All these 13 cases had clinically diagnosed polyneuropathy and in 12 of these 13 cases neuropathy was confirmed in concomitantly analyzed sural nerve biopsies (Fig. 1b). 24 fresh-frozen muscle biopsies from subjects taking part in the MALICoT (Master Athletic Laboratory Study of Intramuscular Connective Tissue) study (registration number DRKS00015764; [19]) served as controls. Of these controls, 9 were 60-75 years old, 6 of which were athletes and 3 unathletic; 15 control subjects were 20-35 years old, of which five were athletes and 10 were unathletic. Available information on gender, age and site of biopsy is included in Supplementary Table 1. Histological findings are reported in Supplementary Table 2-3. On all of these samples, we performed methylation profiling using the Infinium Methylation EPIC BeadChip 850K microarray (Illumina), which tests 850.000 CpG sites in the genome. For comparison, we also included eight cases of multi-minicore myopathy (MMC) from a previously published dataset (GEO accession: GSE121961 [4]) in our analysis. After quality control and data processing, we filtered out unreliable CpG sites, those on sex chromosomes, cross-reactive probes and SNPs. Based on the full post-QC methylation matrix, unsupervised t-SNE (Fig. 1e), principal component analysis (PCA, Supplementary Fig. 1) and the label-free full-matrix sample-to-sample correlation heatmap (Fig. 1f) showed disease-group-associated methylation structure. Reviewer-requested sensitivity analyses showed that the major t-SNE pattern was not explained by a single seed or parameter setting, although t-SNE was interpreted qualitatively because exact two-dimensional cluster geometry varied across settings. PCA sensitivity analyses showed that group-associated structure was also visible beyond one projection or scaling choice; however, leading PCs also overlapped with cohort variables such as sample source, Sentrix ID, age, sex and biopsy site. In these unsupervised analyses, ALS and non-ALS NMA remained partially overlapping. Reviewer-requested subset analyses confirmed that ALS versus non-ALS NMA was the least stable disease-pair comparison, supporting cautious interpretation of this comparison.",
    ),
    87: (
        "Figure 1 legend: describe t-SNE/correlation heatmap cautiously and explicitly retain ALS/NMA overlap.",
        "Figure 1. Disease group-associated and overlapping histological and epigenetic features. Cases were classified based on clinical and histopathological features. Groups of partially atrophic (diameter 20-40 µm) and atrophic (diameter <20 µm) as well as hypertrophic (diameter >80 µm) muscle fibers in neurogenic atrophy due to ALS (a) and sensorimotor neuropathy (b). Arrows in (b): target regions. Cryostat sections, H&E. Scale bars = 50 µm. (c) Marked endomysial inflammatory infiltration with many cytotoxic T cells immunoreactive for CD8 (brown) and focal infiltration of a muscle fiber (arrow) in a case of IBM. Paraffin section, hematoxylin counterstain. Scale bar = 50 µm. (d) Rimmed vacuoles (arrows) in an IBM case. Cryostat sections, H&E. Scale bar = 25 µm. (e) Unsupervised t-SNE visualization of the full post-QC filtered methylation matrix shows disease-group-associated structure, while ALS and non-ALS NMA remain partially overlapping. (f) Label-free sample-to-sample correlation heatmap calculated from the full post-QC M-value matrix and ordered by unsupervised hierarchical clustering; diagnostic and metadata annotations were added only after clustering.",
    ),
    89: (
        "GSEA Results: frame pathway findings as exploratory downstream interpretation.",
        "Next, we performed exploratory gene set enrichment analysis for hypo- and hypermethylated CpG sites in disease groups versus controls (Fig. 2, Supplementary Table 4). Because these analyses are based on disease-group-associated differential methylation in a retrospective pilot cohort, the enriched terms should be interpreted as candidate biological pathways rather than as proof of disease-intrinsic mechanisms. In all disease groups, the top 100 gene ontology (GO) terms associated with both hypo- and hypermethylated CpG sites included (1) those associated with nervous system development, synapse formation, axon morphogenesis – also termed axonogenesis, neurotransmission and function, as well as glial cell development and differentiation; and (2) cell adhesion/ junctions and extracellular matrix. In all disease groups, in particular in the analysis of hypermethylated CpG sites, GO terms associated with muscle development, structure and function as well as cytoskeleton were enriched. In line with the inflammatory nature of these diseases, GO terms associated with inflammation in general and more specifically with T cell response were identified for hypomethylated CpG site in inflammatory myopathies. GO term enrichment associated with myelination/ axon ensheathment was found amongst hypomethylated sites in NMA including ALS versus controls, potentially reflecting disease-associated tissue remodeling and regenerating intramuscular nerve fibers.",
    ),
    91: (
        "GSEA Results continuation: avoid direct upregulation claims from methylation enrichment alone.",
        "Other commonly enriched GO terms included those associated with skeletal muscle development or function. These were particularly enriched for the hypermethylated sites in inflammatory myopathies, but conversely for hypomethylated sites in non-ALS NMA and ALS. Protein degradation-related GO terms, i.e. those associated with autophagy and/or the ubiquitin-proteasome system, were enriched in hypomethylated sites in inflammatory myopathies, non-ALS NMA and ALS, suggesting candidate pathway involvement in these disease groups. In multi-minicore myopathy, these were enriched for the hypermethylated sites, in accordance with the absence of major ongoing muscle fiber degradation usually associated with this disorder.",
    ),
    93: (
        "Figure 2 legend: identify GSEA as downstream differential-methylation-based visualization.",
        "Figure 2. Exploratory gene set enrichment analysis of differentially methylated CpG sites compared to controls. In each comparison, we selected the top differential GO terms and grouped them based on biological/ cellular role (for color-coding see top). For the disease groups (non-IBM-IIM, NOS, IBM, NMA, ALS, MMC), grouped GO term enrichment is shown compared to controls (ctr) in the heatmap based on the p values (hypomethylation on the left, color-coded in different shades of blue; hypermethylation on the right, color-coded in different shades of red). This visualization is based on downstream differential methylation results and is not used as independent unsupervised clustering evidence.",
    ),
    111: (
        "ALS/NMA Results: retain analysis but make it exploratory and avoid implying robust diagnostic separation.",
        "While inflammatory myopathies can in many cases be well differentiated from neurogenic atrophies, including the one occurring in ALS, the distinction within neurogenic atrophy, i.e. ALS NMA versus NMA due to other causes, can be challenging or even impossible solely based on histopathological data. In this pilot cohort, ALS and non-ALS NMA showed partially overlapping methylation profiles, and reviewer-requested subset analyses indicated that this was the least stable disease-pair comparison. We therefore treated the ALS NMA versus non-ALS NMA analysis as exploratory and hypothesis-generating. To explore possible biological correlates of differential methylation within related neurogenic atrophy groups, we performed a CpG site-associated gene set enrichment analysis for the comparison “ALS NMA versus non-ALS NMA”. There were only two GO terms associated with hypermethylated sites, both related to cell adhesion via plasma-membrane adhesion molecules (GO:0007156 and GO:0098742), suggesting possible involvement of cell-adhesion-related methylation differences in this exploratory comparison. Comparing IBM versus non-IBM IIM, NOS revealed that CpG sites associated with genes encoding proteins that are involved in immune system activation, especially T cell/ lymphocyte differentiation/ activation were significantly enriched amongst hypomethylated sites, as were associated with protein degradation and autophagy. Hypermethylated sites in IBM were mostly associated with GO terms linked to the development and differentiation of stem cells/precursor cells, blood vessels and neuronal tissue as well as cell-cell interaction/junctions (Supplementary Table 4).",
    ),
    125: (
        "Results heading: replace diagnosis prediction with disease-group classification.",
        "Supervised learning classifies disease groups in held-out samples within this selected pilot cohort",
    ),
    126: (
        "Supervised learning Results: use leakage-controlled disease-group wording and validation caveat.",
        "Next, we explored whether supervised learning could classify disease groups based on methylation data in this selected pilot cohort. We first split our cases into a training set (60% of the data) and a held-out test set (40%), maintaining the class proportions. Using the training set only, we filtered and preprocessed the data and selected the most discriminative features — removing low-variance CpG sites, then applying univariate and mutual-information selection and Random Forest recursive feature elimination to obtain the 50 most distinct CpG sites. The fitted selectors were applied unchanged to the held-out test set. The training set served to train and tune the individual models (with stratified cross-validation), while the test set was used solely to evaluate the performance of the different supervised learning algorithms. We tested four alternative supervised learning algorithms: logistic regression, decision trees, random forest and SVM (support vector machine). All models classified disease groups in the held-out test set with an accuracy of at least 73% (decision tree, lowest). Logistic regression and random forest performed best, each with an accuracy of 93.3% and a weighted precision of 95.2% (recall 93.3%, F1 92.2%); SVM reached 86.7% accuracy, 92.5% weighted precision, 86.7% recall and F1 86.1%. These performance estimates are internal to this selected cohort and should not be interpreted as validation of a clinical diagnostic classifier. Validation in larger independent cohorts will be required.",
    ),
    128: (
        "Figure 6 legend: avoid true diagnosis/predicted diagnosis wording.",
        "Figure 6. Supervised learning algorithms for disease-group classification. Heatmaps of the correlation matrices are shown for logistic regression (a), decision tree (b), random forest (c) and SVM (d). Logistic regression and random forest showed the highest agreement between clinicopathological disease group and methylation-based disease-group prediction in the held-out test set. Table showing accuracy, precision, recall frequency and F1 score for the different supervised learning models applied to the dataset (e).",
    ),
    132: (
        "Discussion opening: avoid diagnostic-value overclaim and frame as pilot cohort evidence.",
        "In this cohort consisting of inclusion body myositis (IBM), non-IBM inflammatory myopathy (non-IBM IIM, NOS), amyotrophic lateral sclerosis (ALS), non-ALS neuromuscular atrophy (NMA), and multi-minicore myopathy (MMC), as well as normal controls, we observed disease-group-associated CpG methylation patterns that may aid future diagnostic studies and may point at pathophysiological mechanisms in inflammatory and neurogenic atrophy. The diagnostic value for complex cases, mixed pathology, and ALS versus non-ALS NMA requires validation in larger independent cohorts.",
    ),
    133: (
        "Discussion diagnostic tools: soften diagnostic-workup claim.",
        "Methylation profile-based diagnostic tools work particularly well for neoplastic diseases, because these are characterized by expansion of a mutated cell clone. Hence, cells are often more homogenous; minor contamination by the infiltrated host tissue is usually well tolerated. The inhomogeneous nature of healthy and diseased muscle tissue may complicate methylation profiling-based diagnostic tools. Nevertheless, our results suggest that analysis of epigenetic, i.e. DNA methylation patterns, may provide useful complementary information for the study and future diagnostic evaluation of muscle biopsies in selected contexts.",
    ),
    135: (
        "ALS/NMA Discussion: explicitly mark as exploratory and small-cohort limited.",
        "Although not all ALS cases could be distinguished unequivocally from neurogenic atrophy due to other causes, we observed partial separation of several ALS cases from other NMA cases in this pilot study. Because this comparison involved a small number of ALS cases and was sensitive to individual samples, we interpret the ALS versus non-ALS NMA findings as exploratory and hypothesis-generating rather than as evidence of validated diagnostic separation.",
    ),
    136: (
        "Discussion diagnostic application: avoid implying immediate diagnostic-workup use.",
        "In future diagnostic studies, DNA methylation parameters could be evaluated together with histopathological hallmarks of skeletal muscle. In this context, it is interesting to note that fiber type grouping due to collateral reinnervation is often rather weak and only detected in approximately half but not all of the muscle biopsies in ALS [1, 32, 60], whereas NMA in peripheral neuropathy often shows a more chronic course with a more pronounced fiber type grouping. Due to the more rapid onset, muscle biopsies in patients with suspected ALS are frequently taken earlier in the course of the disease than in neurogenic atrophy due to polyneuropathy. Furthermore, unaffected motoneurons in peripheral neuropathies (CMT) have been reported to have a better capacity for reinnervation of adjacent denervated muscle fibers than their counterparts in ALS[60]. In fact, muscle reinnervation in ALS can be impaired. These pathophysiological mechanisms may explain why fiber type grouping is not always present in ALS (detected in approximately half of the cases) and often rather mild, an observation shared by others in comparison to other motor neuron diseases or CMT [1, 32, 60]. Histological distinction of ALS-induced NMA compared to CMT could therefore theoretically be aided by additional molecular features, but this requires validation in larger and clinically balanced cohorts.",
    ),
    138: (
        "ALS biology Discussion: remove diagnostic-role phrasing and keep biological interpretation cautious.",
        "In addition to their possible future diagnostic relevance, exploratory epigenetic differences in ALS muscle may support the hypothesis that skeletal muscle cells in ALS are not only passive bystanders that are simply denervated secondary to motor neuron damage. Hence, independent of muscle denervation, muscle differentiation or energy metabolism may be affected cell-autonomously in ALS, reviewed in [55], which could lead to epigenetic alterations as well. In line with our findings, DNA methylation changes have been identified in spinal cord and skeletal muscle of mouse models of ALS and were found to be therapeutic targets [42, 66].",
    ),
    139: (
        "Discussion pathway interpretation: keep biological interpretation exploratory and cohort-aware.",
        "Our results also suggest candidate common and differential epigenetic pathways in muscle pathophysiology. Gene set enrichment analysis showed involvement of Wnt signaling, in line with its role in regeneration of dystrophic skeletal muscle [26, 27]. While most disease groups showed differential methylation of CpG sites associated with this pathway to some degree, hypermethylation of CpG sites in non-IBM IIM and hypomethylation in multi-minicore myopathy were most notable, suggesting candidate group-associated pathway differences that require validation. Given that in inflammatory myopathies Wnt signaling is activated in muscle fibers, but also systemically [36], the observed hypermethylation of the associated CpG sites could indicate a potentially compensatory mechanism to limit its activation, possibly to limit fibrosis [8]. Furthermore, the correlation of epigenetic with gene expression data in IBM highlighted the roles of epigenetic regulation of the regenerative capacity of muscle satellite stem cells, proteostasis as well as controlling inflammatory signaling in the pathophysiology of this inflammatory myopathy. These findings are in line with and extend previous observations in IBM [29], by pointing at potential molecular mechanisms involved. Because the methylation data were generated from bulk tissue, these pathway interpretations may reflect both disease-associated biology and differences in tissue composition.",
    ),
    140: (
        "Discussion pathway continuation: avoid overinterpreting protein degradation enrichment.",
        "Autophagy dysregulation has been described in IBM rather than in pure PM [61]. Nevertheless, protein degradation-related GO terms, i.e. those associated with autophagy and/or the ubiquitin-proteasome system, were enriched in hypomethylated sites in both IBM and, slightly less pronounced, non-IBM IIM, NOS, further supporting candidate overlap across this inflammatory-myopathy spectrum. However, even in ALS NMA and non-ALS NMA, protein degradation-related GO terms were enriched in hypomethylated sites [3, 15, 30]. In line with these findings, multi-minicore myopathy, a disorder usually not associated with major ongoing muscle fiber degradation, showed hypermethylation of these sites. These pathway findings should be interpreted as exploratory and require validation in larger cohorts with more complete clinical and histopathological annotation.",
    ),
    141: (
        "Conclusion: replace entity-specific profile claim with proof-of-principle disease-group-associated evidence.",
        "In conclusion, the present study provides proof-of-principle evidence that non-neoplastic muscle disorders can show disease-group-associated CpG methylation patterns in a selected cohort of well-defined cases. Furthermore, the observed methylation changes suggest candidate involvement of known pathways in muscle disease pathophysiology, including regulation of stemness, Wnt-signaling, inflammation and altered proteostasis. These findings should be interpreted as exploratory and require validation in larger, independently collected cohorts.",
    ),
    142: (
        "Limitations: add explicit cohort/confounder/ALS validation caveat.",
        "Even in this small cohort, we observed disease-group-associated DNA methylation patterns. However, for the most part, we cannot differentiate between primary changes in muscle cells and secondary changes such as muscle regeneration, inflammatory infiltrates, fibrosis, biopsy-site differences or other tissue-composition effects. In addition, disease group, sample source, Sentrix ID and clinical variables partly overlap in this retrospective pilot cohort and cannot be fully disentangled. For metadata-based sensitivity visualization, we displayed the lymphomonocytic infiltration category only as supplied in the available metadata table and did not independently interpret or re-score this variable. Other validated histopathological severity scores, including fibrosis, necrosis, fiber-type and denervation scores, were not available for all samples and therefore could not be incorporated into the models. Performance estimates of the supervised learning are based on an internal held-out split and may depend on the specific data partitioning and cohort structure. The ALS versus non-ALS NMA comparison is particularly limited by small sample size and individual-sample influence. Hence, validation in larger, independent cohorts using repeated or nested cross-validation will be needed. Furthermore, the application of this method to larger cohorts, the correlation with clinical features such as disease duration, autoantibody status and treatment responses, as well as the inclusion of less well-defined cases and those with dual pathologies will provide more systematic information on its applicability in diagnostic workup and pathophysiological relevance, including the differentiation between primary and secondary changes in the future.",
    ),
}


def replace_with_highlight(paragraph, text: str) -> None:
    paragraph.clear()
    run = paragraph.add_run(text)
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(10)


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document(BASE_DOCX)
    log_lines = [
        "# Manuscript ALS/NMA and supervised-learning wording changes",
        "",
        f"Base manuscript: `{BASE_DOCX}`",
        f"Output manuscript: `{OUT_DOCX}`",
        "",
        "Changed paragraphs are highlighted in yellow in the generated DOCX.",
        "",
    ]

    for idx, (reason, replacement) in REPLACEMENTS.items():
        old = " ".join(doc.paragraphs[idx].text.split())
        if not old:
            raise ValueError(f"Paragraph {idx} is empty; expected text to replace.")
        replace_with_highlight(doc.paragraphs[idx], replacement)
        log_lines.extend(
            [
                f"## Paragraph {idx}",
                "",
                f"Reason: {reason}",
                "",
                "Old text excerpt:",
                "",
                f"> {old[:800]}{'...' if len(old) > 800 else ''}",
                "",
                "Replacement:",
                "",
                replacement,
                "",
            ]
        )

    doc.save(OUT_DOCX)
    CHANGE_LOG.write_text("\n".join(log_lines), encoding="utf-8")
    print(OUT_DOCX)
    print(CHANGE_LOG)


if __name__ == "__main__":
    main()

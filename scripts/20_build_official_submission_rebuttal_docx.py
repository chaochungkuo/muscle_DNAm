#!/usr/bin/env python3
"""Build an official round-2 response-to-reviewers DOCX.

This version is intended for author review before journal submission. It quotes
reviewer comments verbatim by topic and pairs each comment with a direct
response and a concrete manuscript/figure change.
"""

from __future__ import annotations

from pathlib import Path
import shutil

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ANALYSIS_REPO = Path(__file__).resolve().parents[1]
PROJECT_ROOT = ANALYSIS_REPO.parents[1]
REVIEW_DIR = PROJECT_ROOT / "manuscripts" / "reviewer_round_2" / "to_Juliane_2026-07-23"
OFFICIAL_DIR = PROJECT_ROOT / "manuscripts" / "reviewer_round_2" / "official_submission_2026-07-23"
OFFICIAL_FIG_DIR = OFFICIAL_DIR / "figures_response"
SRC_FIG_DIR = REVIEW_DIR / "figures_response"
OUT_DOCX = OFFICIAL_DIR / "Benner_et_al_Response_to_Reviewers_Round2_2026-07-23.docx"
SRC_MANUSCRIPT = REVIEW_DIR / "Bremer_manuscript_clean_highlighted_ALS_NMA_wording_draft.docx"
OUT_MANUSCRIPT = OFFICIAL_DIR / "Benner_et_al_Revised_Manuscript_Highlighted_Round2_2026-07-23.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color: str = "D9D9D9") -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def style_doc(doc: Document) -> None:
    sec = doc.sections[0]
    sec.top_margin = Inches(0.75)
    sec.bottom_margin = Inches(0.75)
    sec.left_margin = Inches(0.75)
    sec.right_margin = Inches(0.75)
    styles = doc.styles
    for style_name in ["Normal", "Title", "Heading 1", "Heading 2", "Heading 3"]:
        styles[style_name].font.name = "Arial"
        styles[style_name]._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    styles["Normal"].font.size = Pt(10)


def add_box(doc: Document, label: str, text: str, fill: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell)
    p = cell.paragraphs[0]
    r = p.add_run(label)
    r.bold = True
    r.font.size = Pt(9)
    p.add_run("\n")
    for i, line in enumerate(text.strip().split("\n")):
        if i:
            p.add_run("\n")
        p.add_run(line.rstrip())
    doc.add_paragraph()


def add_response(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run("Response: ")
    r.bold = True
    p.add_run(text.strip())


def add_changes(doc: Document, text: str) -> None:
    add_box(doc, "Changes made in the manuscript / figures", text, "FFF9E8")


def add_figure(doc: Document, filename: str, caption: str, width: float = 6.5) -> None:
    fig_path = SRC_FIG_DIR / filename
    if not fig_path.exists():
        raise FileNotFoundError(fig_path)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(fig_path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.italic = True
    r.font.size = Pt(8)


def add_section(
    doc: Document,
    heading: str,
    reviewer_text: str,
    response: str,
    changes: str,
    figures: list[tuple[str, str]] | None = None,
) -> None:
    doc.add_heading(heading, level=2)
    add_box(doc, "Reviewer comment (verbatim)", reviewer_text, "F7F7F7")
    add_response(doc, response)
    add_changes(doc, changes)
    for filename, caption in figures or []:
        add_figure(doc, filename, caption)


def build_doc() -> None:
    OFFICIAL_DIR.mkdir(parents=True, exist_ok=True)
    OFFICIAL_FIG_DIR.mkdir(parents=True, exist_ok=True)
    for src in SRC_FIG_DIR.iterdir():
        if src.is_file():
            shutil.copy2(src, OFFICIAL_FIG_DIR / src.name)
    shutil.copy2(SRC_MANUSCRIPT, OUT_MANUSCRIPT)

    doc = Document()
    style_doc(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Response to the Reviewers – Round 2")
    run.bold = True
    run.font.size = Pt(16)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Benner et al.")
    r.font.size = Pt(11)

    add_box(
        doc,
        "General note to the editor and reviewers",
        """
We thank the reviewers for their detailed assessment of the revised manuscript. We have performed additional analyses requested by the reviewers, revised the interpretation of the unsupervised and downstream methylation analyses, softened claims where appropriate, and changed the terminology throughout the manuscript to refer to disease-group-associated methylation structure and disease-group classification rather than diagnostic prediction. All available metadata requested by the reviewers were displayed or analyzed where possible. Variables that were not available in a validated form across the cohort were not inferred.
""",
        "EEF4FF",
    )

    doc.add_heading("Reviewer 2", level=1)
    add_section(
        doc,
        "Reviewer 2, Comment 1 – Inclusion/exclusion criteria and analytical sample set",
        """
The inclusion and exclusion criteria are still not sufficiently detailed. For example, I would be very surprised if the authors' archive included only 6 cases of "non-IBM IIM, NOS" and just 13 cases of non-ALS NMA in the last 8 years. Assuming more cases were available, how were these specific cases selected among all others? Were all candidate cases within the 8-year period reviewed? Were slides reviewed before inclusion, or was inclusion based solely on the initial diagnosis/report? If the latter, how were differences among diagnosing pathologists controlled for? Was clinical information also evaluated, or was inclusion/exclusion based just on the histologic criteria? Etc, etc.
""",
        """
We thank the reviewer for requesting clarification. We revised the manuscript to state what is supported by the analytical record. The methylation analyses included all samples that were provided for this study, fulfilled the stated pre-analytical sample requirements, and passed methylation-array quality control. Samples were not selected or excluded on the basis of methylation patterns, clustering, supervised-learning performance, or any expected analytical result; in other words, there was no analysis-driven filtering of cases. Case groups were defined using the clinicopathological disease-group labels available for the study, and the revised manuscript now avoids implying that the analyzed cohort represents all biopsies or all potential cases from a historical source population.
""",
        """
Methods: clarified that all provided samples meeting the stated tissue/clinical-information requirements and passing methylation-array QC were analyzed. Results: described the final analyzed sample set rather than implying post hoc analytical selection. The revised wording avoids unsupported claims about historical candidate-case totals.
""",
    )

    add_section(
        doc,
        "Reviewer 2, Comment 2 – Tissue composition and deconvolution limitation",
        """
I agree that the longitudinal disease analyses are beyond the scope of the current manuscript. However, that was raised only as an example of a scenario in which a difference in cell-type distribution may have a larger effect on the observed methylation signature than the underlying disease entity. The authors either need to perform deconvolution analyses to address the cell composition concern (by incorporating deconvolution estimates into their differential methylation models) or explicitly discuss that the absence of such analyses is a major limitation of the current study.
""",
        """
We agree that bulk muscle-biopsy methylation profiles can reflect variable tissue and cellular composition, including inflammatory infiltrates, fibrosis, necrosis/regeneration and other biopsy-composition effects. We therefore expanded the manuscript limitation to state that the current bulk-tissue data cannot distinguish primary disease-intrinsic methylation changes from secondary changes related to tissue composition. We displayed the available lymphomonocytic infiltration category only as supplied in the metadata table, without independently interpreting or rescoring it. Validated fibrosis, necrosis, fiber-type and denervation scores were not available for all samples and were therefore not inferred or incorporated into differential methylation models.
""",
        """
Discussion/limitations: added explicit limitation text stating that the absence of complete validated tissue-composition and pathology-severity covariates is a major limitation. Results/Discussion: softened pathway and differential-methylation interpretation accordingly.
""",
    )

    add_section(
        doc,
        "Reviewer 2, Comment 3 – ALS versus non-ALS NMA wording",
        """
Similarly, the ALS/non-ALS comparison has been retained. While that is acceptable, the language in the relevant sections should be further softened to avoid overinterpretation, given the small number of ALS cases.
""",
        """
We agree. We retained the ALS versus non-ALS NMA comparison, but revised the manuscript to describe this analysis as exploratory and hypothesis-generating. The revised text now states that ALS and non-ALS NMA remain partially overlapping in unsupervised analyses, that the ALS versus non-ALS NMA subset was the least stable disease-pair comparison, and that the comparison is sensitive to individual samples. We removed language implying validated diagnostic separation or robust disease-specific classification for ALS versus non-ALS NMA.
""",
        """
Abstract, Results and Discussion: changed ALS/non-ALS NMA wording to exploratory and hypothesis-generating language. Supplementary Response Figure S5 summarizes the ALS/non-ALS NMA robustness analyses.
""",
    )

    add_section(
        doc,
        "Reviewer 2, Comment 4 – Proofreading, disease-group terminology and supplementary table reference",
        """
The revised manuscript appears to have been written in haste and contains many grammatical and typographic errors, as well as instances where revisions have not been fully implemented. To give just a few of many examples:
a. On page 25, a newly added sentence states, "Next, we wanted to find out whether supervised learning can predict a correct diagnosis." Instead of "diagnosis," the correct term should be "disease group." Along the same lines, the abstract states, "Based on the CpG site methylation data, supervised learning, especially using logistic regression and random forest, even predicted diagnosis beyond disease group correctly in many cases." That's an overstatement again, since non-ALS is not a diagnosis, and neither is "non-IBM, NOS" (not to mention that other diseases that would normally be in the differential diagnosis have not yet been evaluated). Similar inconsistencies in replacing the prior "disease-specific" language with more general terms are present elsewhere in the paper, and should be corrected throughout to avoid overstatements.
b. On page 28, the new text refers to "supplementary table xy" instead of "supplementary tables 2 and 3".
The entire manuscript should be carefully proofread before resubmission.
""",
        """
We agree and revised the wording throughout. Supervised learning is now described as disease-group classification within this pilot analytical cohort, not as diagnosis prediction. The abstract, Results, Figure 6 legend and Discussion now state that classifier performance is internal to this cohort and requires validation in independent cohorts. We corrected the supplementary table reference to supplementary tables 2 and 3 and proofread the revised manuscript to remove diagnostic overstatements and inconsistent disease-specific wording.
""",
        """
Global manuscript pass: replaced diagnosis-prediction wording with disease-group classification wording, corrected the supplementary table reference, and removed inconsistent disease-specific or diagnostic overstatements.
""",
    )

    doc.add_heading("Reviewer 1", level=1)
    add_section(
        doc,
        "Reviewer 1, General comment – Study rationale",
        """
This study addresses an important question, whether DNA methylation profiling can reveal reproducible epigenetic profiles across non-neoplastic skeletal muscle diseases. The application of genome-wide methylation analysis to inflammatory, neurogenic, and inherited muscle disorders is relatively novel and could provide useful insights into disease biology and future diagnostic approaches. The assembled cohort and accompanying computational analyses therefore represent a potentially valuable contribution.
""",
        """
We thank the reviewer for this positive assessment of the study rationale and potential value. We agree that the novelty of applying genome-wide DNA methylation profiling to non-neoplastic muscle disease also requires careful attention to cohort structure, tissue composition and technical variables, which we address point by point below.
""",
        """
No specific manuscript change was required for this general positive comment; the subsequent changes address the methodological and interpretative issues raised below.
""",
    )

    add_section(
        doc,
        "Reviewer 1, Comment 1 – Overall concern that unsupervised structure may reflect confounding",
        """
I thank the authors for the clarification that the unsupervised clustering (t-SNE in Figure 1E and PCA in Supplementary Figure 1) was initiated from the full dataset of 771,381 probes without prior diagnosis-based feature selection.

However, this clarification raises concerns regarding the overall methodology, results, and interpretation. The cluster separation observed in the t-SNE and PCA plots is unexpectedly strong given the number of cases and heterogeneous nature of cohort under study. It is essential to establish that this structure is robustly associated with the disease groups rather than technical, demographic, or tissue composition variables before the subsequent supervised analyses and biological interpretations can be evaluated.

My concerns about the unsupervised analysis are substantiated by the following points:

* DNA methylation classification systems have generally been developed in settings where a strong and relatively stable methylation signal is expected. This includes neoplasms, where the methylome reflects cell of origin together with clonal somatic and epigenetic alterations, and Mendelian disorders, where a germline mutation may produce a reproducible epigenetic signature. Even in these settings, published classifiers rely on large reference cohorts on the order of hundreds or thousands of cases. By contrast, the authors explore a more challenging and complex cohort of predominantly acquired, non-neoplastic muscle disease groups, expected to have a much weaker and more heterogeneous epigenetic signature. The authors analyze only 73 cases, including controls, from a small number of diagnostic groups, with as few as six cases per group. In this context, the apparently near-complete cluster separation is unexpected and requires careful demonstration that it is stable and not attributable to the structural distribution of potential confounders.

* The cohort is not balanced or adjusted for many variables that likely have an epigenetic effect in muscle: age, sex (despite removing sex chromosome probes, there are epigenetic sex differences on other chromosomes), biopsy site (individual muscles have different cellular compositions and physiology), comorbidities, disease duration, metabolic/exercise level, diet, and tissue handling. For example, Supplementary table 1 indicates that all controls are soleus biopsies from male adults, whereas all MMCs (a genetic disease) are quadriceps biopsies from young individuals. These confounding variables could contribute both within-group heterogeneity and between-group separation.

* Supplementary Table 1 also suggests that batch effects may be a major confounder. In addition to the external datasets, in-house NMA samples are mostly concentrated on a single chip run. Even though QC and normalization were performed, batch effects can still contribute to the observed cluster separation.

* The analysis also does not adequately account for biopsy cellular composition: depending on the underlying pathology, there may be varying contributions from adipose tissue, connective tissue, inflammatory cells, different fiber types, necrosis/regeneration, and fibrosis. The authors acknowledged that using bulk tissue is a limitation in their rebuttal to Reviewer 2. Nevertheless, depending on how patchy the disease process is, the measured methylation profile may primarily reflect differences in tissue or cellular composition or the proportion of histologically affected tissue rather than disease related alterations.
""",
        """
We thank the reviewer for raising these important points. We agree that bulk DNA methylation profiles from muscle biopsies are influenced by tissue composition, including variable contributions from muscle fibers, adipose tissue, connective tissue, inflammatory infiltrates, necrotic or regenerating fibers, and fibrosis. We also agree that sample source, Sentrix array, biopsy site, age and sex partly overlap with disease group in this retrospective cohort. Other potentially relevant variables highlighted by the reviewer, including comorbidities, disease duration, metabolic or exercise level, diet and tissue handling, were not available in a complete and validated form across the cohort and therefore could not be modeled. We therefore revised the interpretation throughout the manuscript. The additional analyses support disease-group-associated methylation structure, while showing that this structure cannot be interpreted as a purely disease-intrinsic or cell-type-specific methylation signature. Instead, the observed profiles likely reflect the composite methylation phenotype of routine diagnostic muscle biopsies. We now present this work as exploratory and hypothesis-generating, requiring validation in larger, prospectively collected and independently processed cohorts.
""",
        """
Results/Discussion: revised wording from disease-intrinsic or diagnostic-signature language to disease-group-associated methylation structure. Limitations: added that disease group, sample source, Sentrix ID, biopsy site and tissue-composition variables are structurally linked in this retrospective pilot cohort and cannot be fully disentangled.
""",
        [
            (
                "Response_Figure_1_unsupervised_structure.png",
                "Response Figure 1. Unsupervised methylation structure and cohort variables. Baseline PCA, t-SNE, PC-metadata associations and full-matrix correlation heatmap show disease-group-associated structure while documenting overlap with cohort variables.",
            )
        ],
    )

    add_section(
        doc,
        "Reviewer 1, Comment 2 – t-SNE default PCA, pca=FALSE and parameter sensitivity",
        """
The authors state in the rebuttal that they used the entire set of 771,381 probes for the t-SNE visualization using Rtsne(t(getM(mSetSq)), perplexity = 15, theta = 0.5, dims = 2). However, according to the official documentation for Rtsne, default parameters include pca = TRUE and initial_dims = 50. Thus, although the input comprises all probes and the procedure remains unsupervised, the t-SNE is actually calculated from the first 50 internally derived principal components rather than directly from the complete probe space. Those components may themselves be dominated by the technical and biological variables described above. This initial PCA step is standard and computationally reasonable, but it should be reported explicitly. Please include a direct analysis with pca = FALSE if computationally feasible. Please provide a systematic sensitivity analysis using several feasible values of initial_dims (for example, 10, 20, 30, 50, and 72), perplexities (for example, 5, 10, 15, and 20), and random seeds. Where possible, clustering stability should also be summarized quantitatively rather than assessed solely by visual inspection.
""",
        """
We thank the reviewer for pointing out the Rtsne default. We now state explicitly that the baseline t-SNE used pca=TRUE and initial_dims=50, meaning that Rtsne first reduced the input matrix to 50 internally derived principal components before computing the two-dimensional t-SNE embedding. To test whether the visualization depended on this default or on a single arbitrary parameter setting, we repeated t-SNE across initial_dims values of 10, 20, 30, 50 and 72, perplexities of 5, 10, 15 and 20, and multiple random seeds, and also performed a direct full-matrix t-SNE with pca=FALSE. Stability was summarized quantitatively using Procrustes similarity and silhouette values in addition to visual inspection. These sensitivity analyses show that the major separation pattern is reproducible across a range of t-SNE settings and is therefore not explained by a single random seed or one arbitrary parameter choice. However, the exact position and distance between clusters vary between runs, as expected for t-SNE. We therefore use t-SNE as a qualitative visualization of group-associated structure, not as a quantitative measure of global distances between disease groups.
""",
        """
Methods: added explicit Rtsne parameters, including pca=TRUE and initial_dims=50 for the baseline analysis, and the direct pca=FALSE sensitivity analysis. Results/Discussion: added sensitivity-analysis interpretation and avoided quantitative interpretation of t-SNE distances.
""",
        [
            (
                "Response_Figure_2_tsne_pca_sensitivity.png",
                "Response Figure 2. t-SNE and PCA sensitivity analyses. The major structure persists across t-SNE settings and PCA scaling choices, but t-SNE geometry remains parameter-dependent and should be interpreted qualitatively.",
            )
        ],
    )

    add_section(
        doc,
        "Reviewer 1, Comment 3 – Metadata coloring and unavailable covariates",
        """
For the primary t-SNE analysis and the sensitivity analyses above, please display the same coordinates with samples colored separately by potential confounder variables: dataset source (MALICoT controls, GEO, and the institutional archive), batch information (chip/array/Sentrix), age, sex, biopsy site, estimated inflammatory cell fraction (lymphocytes/macrophages), fiber type estimates, denervation estimates and pathology severity using an appropriate fibrosis/necrosis score. These annotations should be provided where variables are available or can be estimated using validated methods and would help determine whether the confounding variables explain or contribute significantly to the clustering.
""",
        """
As suggested, we now provide PCA and t-SNE plots colored by available metadata variables, including disease group, sample source, Sentrix ID, age, sex, biopsy-site group and lymphomonocytic infiltration category where provided. The lymphomonocytic infiltration category is displayed only as supplied in the metadata table; we do not independently interpret the scoring scheme or convert uncertain annotations into low/intermediate/high categories ourselves. Requested variables not available in the provided metadata table, including validated fibrosis, necrosis, fiber-type and denervation scores, were not inferred or modeled. These visualizations show that several cohort variables overlap with the disease-group structure, which is now explicitly reflected in the revised interpretation.
""",
        """
Supplementary Response Figures S1 and S2 show PCA and t-SNE colored by available metadata. The manuscript and rebuttal now distinguish available variables from unavailable variables and do not infer unvalidated histopathological scores.
""",
    )

    add_section(
        doc,
        "Reviewer 1, Comment 4 – Label-free full-matrix correlation heatmap and Figure 1F",
        """
Please provide a sample-to-sample correlation heatmap calculated from the full post-QC M-value matrix to determine whether the separation structure is visible in the original methylation data, since t-SNE prioritizes local neighborhood structure and does not preserve global distances. Samples should be hierarchically clustered using a prespecified correlation-based distance without using diagnostic labels to determine their ordering. Diagnosis and potential confounders should be added only as annotations at the end of the analysis.

The heatmap in Figure 1F is misleading when shown together with the t-SNE and presented in the first paragraph of the Results. According to the code referenced in the manuscript, probes were selected using the diagnostic labels:
# Create a contrast matrix for "one vs control" comparisons
contMatrix <- makeContrasts(
ALS_vs_Control = ALS - Control,
IBM_vs_Control = IBM - Control,
Multiminicores_vs_Control = Multiminicores – Control,
NMA_vs_Control = NMA - Control,
PM_vs_Control = PM - Control,
ALS_vs_NMA = ALS - NMA,
IBM_vs_PM = IBM - PM,
levels = design)
The authors then selected the top 500 ranked CpGs from each of these seven comparisons and plotted their union. This is a supervised analysis, and separation by diagnostic group is partly expected by construction and should not be presented as independent evidence that the samples cluster according to disease. If the authors meant to show the heatmap as a supervised descriptive visualization only, that should be clearly stated in the text and figure legend. If the heatmap was meant to further support the unsupervised clustering of the disease groups, it should be replaced or supplemented by the unsupervised sample-to-sample correlation heatmap described above.
""",
        """
We agree that the previous top-CpG heatmap was label-informed and should not be presented as independent unsupervised evidence. We replaced it with a sample-to-sample Pearson correlation heatmap calculated from the complete post-QC M-value matrix. Samples were ordered by hierarchical clustering using 1 minus Pearson correlation and complete linkage, and diagnostic and metadata annotations were added only after clustering. The previous label-informed top-CpG heatmap, if retained, will be described only as a supervised descriptive visualization and not as evidence of unsupervised clustering.
""",
        """
Figure 1F: replaced with a full-matrix label-free sample-to-sample correlation heatmap and added an annotation legend. Manuscript text/legend: distinguishes the label-free full-matrix heatmap from any label-informed top-CpG descriptive heatmap.
""",
    )

    add_section(
        doc,
        "Reviewer 1, Comment 5 – PCA evaluation, scaling and PC-metadata association",
        """
The PCA analysis also requires more detailed evaluation. According to the available code, PCA was performed using prcomp(t(getM(mSetSq)), scale.=TRUE). Please provide a scree plot reporting the variance explained by at least the first 20 principal components together with cumulative variance explained, and show multiple pairwise projections, including PC1-PC2, PC1-PC3, PC2-PC3, and subsequent components where relevant. Evaluation limited to PC1 and PC2 may overlook technical or biological structure represented by other major components. The PCA coordinates should be displayed with samples colored separately by diagnostic group, dataset source, Sentrix/chip, biopsy site, sex, age, and available tissue-composition or histopathologic estimates, as requested for the t-SNE above. In addition to visual inspection, please quantify the association of each of the leading PCs with these variables,. This would help determine whether the variance represented by the principal components is more strongly associated with diagnosis or with technical and demographic structure. Because the analysis uses scale.=TRUE, each CpG is standardized to unit variance before PCA. Scaling hundreds of thousands of probes may give low-variance or noisy probes equal weighting with biologically variable probes. Please explain the rationale for the primary scaling approach, and provide an analysis using centered but unscaled M-values (scale.=FALSE). The authors should also examine the CpGs with the largest positive and negative loadings on the leading PCs to assess whether those components are associated with disease biology, batch, dataset source, age, sex, biopsy site, or tissue composition.
""",
        """
We expanded the PCA analysis in four ways. First, we added scree and cumulative-variance plots for the first 20 PCs to show how much methylation variance is captured by the leading components. Second, we plotted PC1-PC2, PC1-PC3 and PC2-PC3 to test whether the observed structure is restricted to the first two PCs or also appears in other major components. Third, we annotated the strongest positive and negative CpG loadings of the leading PCs to support interpretation of whether major components reflected disease group, cohort variables or technical structure. Fourth, we repeated PCA using centered but unscaled M-values to test whether the observed structure depends on unit-variance scaling of CpGs. These analyses show that disease-group-associated structure is present in PCA, but that leading PCs also overlap with cohort variables such as source, Sentrix ID, age, sex and biopsy site. We therefore retained the PCA result but revised the interpretation to avoid presenting the PCs as disease-only axes.
""",
        """
Methods: added scaled PCA (scale.=TRUE) and unscaled PCA sensitivity analysis (scale.=FALSE). Results: added scree/cumulative variance, PC-metadata association and loading interpretation. Supplementary Response Figure S3 provides the full PCA sensitivity figure.
""",
    )

    add_section(
        doc,
        "Reviewer 1, Comment 6 – Recomputed subset analyses and influential samples",
        """
Re-run the t-SNE and PCA in subsets that remove the dominant or externally sourced groups and directly test the clinically relevant distinctions: (1) excluding MMC; (2) excluding controls; (3) institutional archive samples only; (4) IBM versus non-IBM IIM; and (5) ALS versus non-ALS NMA. These analyses should be recomputed within each subset rather than merely removing points from an embedding calculated from the complete cohort. Given the small group sizes, the authors should also assess whether the observed structure is sensitive to individual influential samples.
""",
        """
We recomputed PCA and t-SNE within each requested subset: excluding MMC, excluding controls, in-house data, IBM versus non-IBM IIM, and ALS versus non-ALS NMA. These analyses test whether the full-cohort structure is driven by one dominant group or by inclusion of external/public data. We use the term in-house data as an analysis label for the non-public/non-GEO methylation samples generated within this study; we do not use it to imply a single-source or otherwise homogeneous cohort. The subset analyses showed that the full-cohort pattern was stronger than several clinically focused subsets, and ALS versus non-ALS NMA was the least stable comparison. We therefore retained this comparison only as exploratory and hypothesis-generating.
""",
        """
Supplementary Response Figure S4 provides independently recomputed subset PCA/t-SNE analyses. Supplementary Response Figure S5 focuses on ALS versus non-ALS NMA robustness and individual-sample influence. Results/Discussion: ALS versus non-ALS NMA is described as exploratory and sample-sensitive.
""",
    )

    add_section(
        doc,
        "Reviewer 1, Comment 7 – Downstream analyses after confounding assessment",
        """
The unsupervised analyses should be resolved before the downstream disease-specific claims can be interpreted confidently. The manuscript and letter of rebuttal appear to use the PCA and t-SNE findings as evidence that the methylation data contain an intrinsic structure corresponding to the diagnostic groups and then proceeds to identify differentially methylated CpGs, construct label-informed heatmaps and classifiers, and assign biological meaning to the resulting genes and pathways. If the apparent unsupervised separation is instead substantially driven by confounders, these same variables would also be expected to influence the subsequent supervised and differential analyses. Therefore, confirmation that the unsupervised structure is robust to these potential confounders is an important prerequisite for interpreting the downstream results as associated to disease groups. Even if the unsupervised findings are confirmed, the differential methylation and supervised-learning analyses will still require their own confounder validation.
""",
        """
The additional analyses support the presence of disease-group-associated methylation structure, while also showing that cohort variables such as source, Sentrix ID, age, sex and biopsy site partly overlap with disease group in this retrospective pilot cohort. Where statistically possible, we fitted differential methylation models adjusted for age/sex or biopsy site and found that several disease-group-associated results persisted, whereas the ALS versus non-ALS NMA contrast was more sensitive to adjustment. We also rebuilt supervised learning using patient-aware splitting and training-only feature selection, and added metadata-only classification as a confounding diagnostic. These analyses do not negate the disease-group-associated findings; rather, they define their appropriate scope. We retained the differential methylation, gene set enrichment and supervised-learning findings, but revised the manuscript to avoid overclaiming clinical diagnostic performance, disease-intrinsic specificity or final mechanistic proof from pathway enrichment.
""",
        """
Results/Discussion: supervised learning is presented as leakage-controlled disease-group classification within this pilot analytical cohort, not as a validated clinical diagnostic classifier. Differential methylation and gene set enrichment are interpreted using exploratory and candidate-pathway language. Response Figure 3 summarizes downstream robustness checks.
""",
        [
            (
                "Response_Figure_3_downstream_robustness.png",
                "Response Figure 3. Downstream robustness checks, including covariate-adjusted differential methylation where estimable, metadata-only classification and patient-aware supervised learning. ALS/non-ALS NMA subset and influence analyses are provided separately in Supplementary Response Figure S5.",
            )
        ],
    )

    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_heading("Supplementary response figures provided with this response", level=1)
    for item in [
        "Supplementary Response Figure S1: PCA colored by available metadata variables.",
        "Supplementary Response Figure S2: t-SNE colored by available metadata variables.",
        "Supplementary Response Figure S3: PCA sensitivity analyses, including scree/cumulative variance, multiple PC projections and unscaled PCA.",
        "Supplementary Response Figure S4: subset PCA/t-SNE analyses recomputed within each requested subset.",
        "Supplementary Response Figure S5: exploratory ALS versus non-ALS NMA robustness and individual-sample influence analyses.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.save(OUT_DOCX)
    print(OUT_DOCX)
    print(OUT_MANUSCRIPT)


if __name__ == "__main__":
    build_doc()

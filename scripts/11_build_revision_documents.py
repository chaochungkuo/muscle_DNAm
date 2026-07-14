from pathlib import Path
import tempfile, zipfile
import xml.etree.ElementTree as ET
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Pt

ROOT=Path(__file__).resolve().parents[1]
PROJECT=ROOT.parents[1]
MANUSCRIPTS=PROJECT/'manuscripts'
OUT=MANUSCRIPTS/'reviewer_round_2/final_drafts'; OUT.mkdir(parents=True,exist_ok=True)
SRC=MANUSCRIPTS/'2026 06 02_Bremer_manuscript (1).docx'

def accepted_changes_copy(src, dst):
    ns='http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    with zipfile.ZipFile(src) as zin, zipfile.ZipFile(dst,'w',zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data=zin.read(item.filename)
            if item.filename=='word/document.xml':
                root=ET.fromstring(data)
                # Remove deleted/moved-from content; unwrap inserted/moved-to content.
                for parent in root.iter():
                    for child in list(parent):
                        tag=child.tag.rsplit('}',1)[-1]
                        if tag in {'del','moveFrom'}:
                            parent.remove(child)
                        elif tag in {'ins','moveTo'}:
                            idx=list(parent).index(child)
                            parent.remove(child)
                            for grand in list(child):
                                parent.insert(idx,grand); idx+=1
                data=ET.tostring(root,encoding='utf-8',xml_declaration=True)
            zout.writestr(item,data)

def find(doc,start):
    for p in doc.paragraphs:
        if p.text.strip().startswith(start): return p
    raise KeyError(start)

def replace(doc,start,text):
    p=find(doc,start); p.text=text; return p

def replace_inside(doc,old,new):
    for p in doc.paragraphs:
        if old in p.text:
            p.text=p.text.replace(old,new); return p
    raise KeyError(old)

accepted=OUT/'_accepted_changes_working_copy.docx'
accepted_changes_copy(SRC,accepted)
doc=Document(accepted)
replace(doc,"Disease group-specific DNA methylation patterns", "Disease group-associated DNA methylation patterns")
replace(doc,"Here we performed a cross-sectional study", "Here we performed a cross-sectional pilot study using the Illumina EPIC methylation array to characterize group-associated CpG methylation patterns in inflammatory myopathies (inclusion body myositis, IBM, and non-IBM idiopathic inflammatory myopathy, not otherwise specified), neurogenic atrophy (amyotrophic lateral sclerosis, ALS, and non-ALS cases), multi-minicore myopathy, and control skeletal muscle.")
replace(doc,"T-SNE analysis and hierarchical clustering", "Unsupervised PCA, t-SNE and sample-correlation analyses demonstrated strong structure associated with the studied disease groups. However, sensitivity analyses also identified substantial associations with dataset source, Sentrix array, age, sex and biopsy site. Patient-aware supervised learning classified the studied groups with high held-out accuracy, but metadata alone also predicted group membership, indicating that the present pilot cohort cannot establish a clinically validated or disease-intrinsic classifier. Differential-methylation sensitivity analyses supported several group-associated biological processes but showed that some results, particularly ALS versus non-ALS neurogenic atrophy, were sensitive to biopsy-site adjustment. Gene-set analyses implicated cytoskeletal maintenance, cell adhesion, muscle and neural development, Wnt signaling, proteostasis and immune-cell-related processes. Correlation with IBM expression data suggested candidate epigenetic mechanisms involving muscle stem-cell activation and regeneration.")
replace(doc,"This study shows that CpG site methylation profiling", "This pilot study shows that bulk-muscle CpG methylation profiling contains reproducible structure associated with the studied pathological groups and can generate pathophysiological hypotheses. Larger, prospectively balanced and independently validated cohorts are required to separate disease-associated signals from technical, demographic, biopsy-site and cellular-composition effects.")
replace(doc,"Non-neoplastic muscle disorders are a heterogeneous", "Non-neoplastic muscle disorders are heterogeneous and include genetic, metabolic, toxic, inflammatory and neurogenic conditions. Although epigenetic alterations have been reported in selected muscle diseases, robust disease-entity-specific methylation signatures have not been established. We therefore investigated whether bulk-muscle methylation contains patterns associated with broad pathological disease groups and whether these patterns highlight candidate biological mechanisms. Machine learning was used as an exploratory classification analysis within the studied groups, not as a clinically validated diagnostic test.")

p=find(doc,"Fresh frozen human skeletal muscle biopsy material")
note=p.insert_paragraph_before()
r=note.add_run("JULIANE DECISION NEEDED: Add the total number of archive candidates screened for each group; confirm whether all candidates within the eight-year period were reviewed; describe slide re-review, clinical-information review, diagnosing pathologists/adjudication, and exact exclusion counts/reasons. The bioinformatics files cannot establish these facts.")
r.font.highlight_color=WD_COLOR_INDEX.YELLOW; r.bold=True

replace(doc,"Array data analysis was performed using R v.4.3.3", "The original array preprocessing was performed in R v4.3.3. Raw IDAT intensities were read with minfi, quantile normalized, and filtered by detection P value, sex-chromosome location, nearby SNPs and cross-reactivity, leaving 771,381 CpG probes. M values were used for statistical analysis. The reviewer-round reanalysis was reproduced in a Pixi-locked R v4.4.3/Python 3.11 environment. Scaled PCA used prcomp with centering and unit-variance scaling; a centered, unscaled PCA was added as a sensitivity analysis. Baseline t-SNE used Rtsne with perplexity 15, theta 0.5, seed 42 and its default internal centered, unscaled PCA (pca=TRUE, initial_dims=50). Sensitivity analyses varied initial_dims (10, 20, 30, 50 and 72), perplexity (5, 10, 15 and 20), and ten random seeds; a direct full-matrix analysis with pca=FALSE was also performed. The same coordinates were annotated separately by disease group, dataset source, Sentrix ID, age, sex, biopsy site, city of origin and the supplied lymphomonocyte category. The latter was not used as a covariate because its scoring provenance was unavailable to the analyst. Leading-PC associations were quantified by eta-squared, and top positive/negative loadings were annotated. A label-free sample-correlation heatmap used 1 minus Pearson correlation and complete-linkage clustering, with labels added only after clustering. PCA and t-SNE were recomputed independently after excluding MMC, excluding controls, restricting to the institutional archive, comparing IBM with non-IBM IIM, NOS, and comparing ALS with non-ALS NMA. Differential methylation used limma with Benjamini-Hochberg FDR correction. Age/sex and biopsy-site sensitivity models were fitted when estimable; source and Sentrix models were retained as design audits when rank deficiency prevented separation from disease group. The cohort comprised 73 samples from 72 patients; two non-IBM IIM samples came from one patient. Because only one patient was repeated, a common within-patient correlation was not estimable, and sensitivity analyses instead retained either one of the paired samples.")
replace(doc,"For supervised classification, the dataset was first split", "For exploratory supervised classification, samples were partitioned at patient level so that the two samples from one patient could not cross training and held-out test sets. The final partition contained 44 training samples from 43 patients and 29 held-out samples from 29 patients. All feature selection was fitted using training samples only. Starting from 771,381 QC-filtered CpGs, training-set variance filtering, ANOVA F-test selection, mutual-information selection, scaling and random-forest recursive feature elimination retained 50 CpGs. Fitted transformations were applied unchanged to the held-out samples.")
replace(doc,"Model hyperparameters were tuned", "Hyperparameters were tuned by three-fold stratified group cross-validation within the training cohort, using patient ID as the grouping variable. Logistic regression, decision tree, random forest and support vector machine models were implemented in scikit-learn. The held-out test cohort was used once for final evaluation. Metadata-only classifiers were evaluated by repeated stratified cross-validation to quantify disease-group information contained in source, Sentrix, age, sex and biopsy site.")
replace(doc,"Supervised Learning for Predicting Diagnosis", "Exploratory Supervised Classification of the Studied Disease Groups")

replace(doc,"Disease groups cluster together in PCA", "Disease-group-associated methylation structure is accompanied by technical and cohort structure")
replace(doc,"When performing hierarchical clustering considering only top 10", "Unsupervised PCA and t-SNE showed strong structure associated with the studied groups, although ALS and non-ALS NMA overlapped. The exact t-SNE geometry varied across initial dimensions, perplexities and random seeds; a direct full-matrix t-SNE with pca=FALSE was feasible and provided an additional sensitivity analysis. Leading PCs were associated not only with disease group but also with dataset source, Sentrix ID and other supplied metadata. A label-free full-matrix sample-correlation heatmap was therefore added as the primary unsupervised hierarchical visualization. The previous top-CpG heatmap was generated from limma-ranked, label-informed CpGs and is presented only as a supervised descriptive visualization, not independent evidence of natural clustering.")
replace(doc,"Figure 1. Disease group-specific", "Figure 1. Histological findings and unsupervised methylation structure. Cases were classified using clinical and histopathological information. (A-D) Representative histological findings as described previously. (E) Baseline unsupervised t-SNE of all 771,381 post-QC CpGs; Rtsne internally reduced the input to 50 centered, unscaled principal components before t-SNE. (F) Label-free sample-to-sample Pearson-correlation heatmap calculated from the complete post-QC M-value matrix and hierarchically clustered using 1 minus Pearson correlation and complete linkage. Disease group and potential confounders were added only as annotations and did not determine sample ordering. Parameter, confounder-coloring and subset sensitivity analyses are shown in the Supplementary Figures.")

replace(doc,"Direct comparison between related disease entities", "Exploratory comparison between related disease groups")
replace(doc,"While inflammatory myopathies can in many cases", "The distinction between ALS-associated and other neurogenic atrophy can be difficult using histopathology alone. In this small cohort, the unadjusted ALS-versus-NMA analysis yielded only two CpGs at FDR <0.001; age/sex adjustment yielded three, whereas biopsy-site adjustment yielded none. Correspondingly, ALS and non-ALS NMA overlapped in unsupervised analyses and the result was sensitive to individual samples. The ALS-versus-NMA analysis is therefore exploratory and hypothesis-generating. Any associated gene-set findings, including cell-adhesion terms, require confirmation in a larger, balanced cohort. The IBM-versus-non-IBM IIM, NOS comparison likewise remains exploratory and may reflect inflammatory-cell composition and disease duration in addition to within-muscle epigenetic regulation.")

replace(doc,"Supervised learning can predict diagnosis", "Exploratory supervised learning classifies the studied disease groups")
replace(doc,"Next, we wanted to find out whether supervised learning", "We evaluated whether methylation data could classify samples among the six studied groups while keeping samples from the same patient in one partition. The patient-aware split contained 44 training samples from 43 patients and 29 held-out samples from 29 patients. Feature selection was performed on training samples only and retained 50 CpGs. Logistic regression achieved 96.6% held-out accuracy and 91.7% balanced accuracy; decision tree achieved 93.1% and 88.3%; SVM achieved 93.1% and 83.3%; and random forest achieved 89.7% and 80.0%, respectively. These values describe classification within this selected cohort rather than clinical diagnostic performance. In repeated cross-validation, defined metadata alone predicted disease group substantially above chance, and Sentrix ID alone was also informative. Because dataset source and Sentrix were structurally aligned with some disease groups, source- or array-independent generalization could not be established.")
replace(doc,"Figure 6. Supervised learning algorithms", "Figure 6. Patient-aware exploratory classification of the studied disease groups. Samples from the same patient were restricted to one partition. Bars show held-out accuracy, balanced accuracy and weighted F1 for logistic regression, decision tree, random forest and SVM after training-only feature selection and patient-group-aware cross-validation. Results describe internal classification of this selected pilot cohort and are not an externally validated clinical diagnostic test.")

replace(doc,"In this cohort consisting of inclusion body", "In this pilot cohort, bulk-muscle CpG methylation showed strong structure associated with inflammatory myopathy, neurogenic atrophy, multi-minicore myopathy and control groups. However, disease group, dataset source, Sentrix array, age, sex and biopsy site were imbalanced and partly aligned. Consequently, the observed patterns cannot be attributed exclusively to disease-intrinsic methylation, and their clinical diagnostic value remains undetermined.")
replace(doc,"Methylation profile-based diagnostic tools work", "Methylation classifiers are most established in settings with strong and stable signals and large reference cohorts. Bulk diseased muscle is heterogeneous and contains variable muscle fibers, inflammatory cells, fibroblasts, vessels, adipose tissue, necrosis, regeneration and fibrosis. Sample-level deconvolution estimates suitable for covariate adjustment were not available to the analyst for incorporation into the present differential models. The absence of cell-composition-adjusted analyses is therefore a major limitation. The current results should be regarded as proof-of-concept group-associated structure rather than a diagnostic tool.")
replace(doc,"Although not all cases of ALS", "ALS and non-ALS NMA were not unequivocally separated. The very small ALS group, individual-sample influence and loss of FDR-significant CpGs after biopsy-site adjustment require that this comparison be considered preliminary and hypothesis-generating.")
replace(doc,"In conclusion, the present study demonstrates", "In conclusion, this pilot study demonstrates that bulk-muscle methylation contains strong structure associated with the studied pathological groups and highlights candidate biological pathways. The study does not establish disease-entity-specific methylation signatures or a clinically validated diagnostic classifier because technical, demographic, biopsy-site and cellular-composition effects cannot be fully separated in this cohort.")
replace(doc,"Even in this small cohort", "The findings require validation in larger, prospectively collected, balanced and independent cohorts. Such studies should include standardized biopsy sites and tissue handling, patient-level grouping, cell-composition or histopathology covariates, clinically relevant differential diagnoses, and repeated or nested validation. The present supervised estimates are internal to one selected cohort and should not be generalized to routine diagnosis.")
replace(doc,"he application of this method", "Application to larger cohorts, correlation with disease duration, autoantibodies and treatment, and inclusion of mixed or diagnostically difficult cases will be necessary to determine whether any methylation features add value beyond histopathology and clinical information.")
replace_inside(doc,"Histological features are now included in supplementary table xy", "Histological features are included in Supplementary Tables 2 and 3")
replace_inside(doc,"Supplementary Table XY", "Supplementary Tables 2 and 3")

# Add a visible draft banner without altering the original file.
banner=doc.paragraphs[0].insert_paragraph_before()
run=banner.add_run("CLEAN REVISION DRAFT — ANALYTICAL CHANGES IMPLEMENTED; JULIANE DECISION ITEMS HIGHLIGHTED")
run.bold=True; run.font.highlight_color=WD_COLOR_INDEX.YELLOW; run.font.size=Pt(12)
doc.save(OUT/'Bremer_manuscript_reviewer_round_2_clean_draft.docx')
accepted.unlink(missing_ok=True)

responses = [
("Reviewer 1 — overall concern regarding confounding", "We agree. We rebuilt the analysis in a locked, reproducible environment and systematically evaluated disease group together with dataset source, Sentrix ID, age, sex, biopsy site, city and the supplied lymphomonocyte category. Disease group was strongly associated with PC1, but technical and cohort variables were also associated with leading PCs. We now state that the observed structure is group-associated and cannot be attributed exclusively to disease-intrinsic methylation."),
("Rtsne internal PCA and sensitivity", "We now explicitly report that baseline Rtsne used pca=TRUE, centered/unscaled internal PCA and initial_dims=50. We evaluated initial_dims 10/20/30/50/72, perplexity 5/10/15/20 and ten seeds (200 runs), quantified Procrustes similarity and silhouette, and performed the requested direct pca=FALSE analysis on all 771,381 probes. Exact two-dimensional geometry varied, so cluster placement is no longer overinterpreted."),
("Coloring by potential confounders", "The same PCA/t-SNE coordinates are now displayed separately by disease group, dataset source, Sentrix, age, sex, biopsy-site group, city and the supplied lymphomonocyte category. The lymphomonocyte score is not interpreted or modeled because its scoring provenance was unavailable to the analyst. Other requested histopathology variables were not supplied; this is identified for author decision and limitation wording."),
("Sample-correlation heatmap", "We added a label-free Pearson sample-correlation heatmap from the complete post-QC M-value matrix, clustered with 1-correlation and complete linkage. Labels were added only after clustering. This replaces the label-informed heatmap as the primary unsupervised hierarchical display."),
("PCA scrutiny", "We added PC1-PC20 scree and cumulative-variance plots, PC1-PC2/PC1-PC3/PC2-PC3 projections, centered-scaled and centered-unscaled PCA, quantitative PC-metadata associations, and annotated the strongest positive and negative loadings for PCs 1-5."),
("Subset and influential-sample analyses", "PCA and t-SNE were recomputed independently after excluding MMC, excluding controls, restricting to the institutional archive, IBM versus non-IBM IIM, NOS, and ALS versus non-ALS NMA. Individual-sample influence was quantified for the two clinically relevant contrasts. ALS-NMA remains unstable and is now explicitly exploratory."),
("Figure 1F", "We agree that the previous top-CpG heatmap was label-informed. It is now described only as supervised descriptive visualization and is moved/relegated to supplementary material. The label-free correlation heatmap is proposed for Figure 1F."),
("Downstream confounder validation", "Age/sex and biopsy-site limma sensitivity models were estimable. Source and Sentrix models were rank deficient with disease group and are reported as non-estimable. ALS-NMA had 2 FDR<0.001 CpGs unadjusted, 3 after age/sex adjustment and 0 after biopsy-site adjustment. Metadata-only classification was also high, demonstrating that classifier performance cannot be interpreted as disease-intrinsic or clinical validation."),
("Reviewer 2 — case inclusion and exclusion", "JULIANE DECISION NEEDED: provide archive candidate totals, whether all cases in the eight-year interval were reviewed, slide and clinical re-review procedure, pathologist adjudication, and exact exclusion counts/reasons. We have not invented information unavailable to the analyst."),
("Reviewer 2 — cell composition", "We now describe the absence of cell-composition-adjusted differential models as a major limitation. The supplied lymphomonocyte categories were not modeled because their definition was unavailable. JULIANE DECISION NEEDED: confirm whether validated deconvolution fractions or histopathology scores can be supplied and incorporated."),
("Reviewer 2 — ALS/non-ALS language", "All ALS-NMA statements have been softened to exploratory/hypothesis-generating. We report the small sample size, overlap, individual influence and biopsy-site sensitivity."),
("Reviewer 2 — terminology and proofreading", "We replaced diagnosis/entity-specific claims with disease-group-associated or classification among the studied groups, changed PM display terminology to non-IBM IIM, NOS, corrected Supplementary Tables 2 and 3 placeholders, and prepared a clean language-revised manuscript draft."),
]

md=["# Point-by-point rebuttal draft","","Analytical responses are complete. Items requiring clinical/pathology knowledge are explicitly marked for Juliane.",""]
rdoc=Document(); rdoc.add_heading('Point-by-point rebuttal draft',0)
for head,body in responses:
    md += [f"## {head}","",body,""]
    rdoc.add_heading(head,level=1); p=rdoc.add_paragraph(body)
    if body.startswith('JULIANE'): p.runs[0].font.highlight_color=WD_COLOR_INDEX.YELLOW
(OUT/'Point_by_point_rebuttal_draft.md').write_text('\n'.join(md))
rdoc.save(OUT/'Point_by_point_rebuttal_draft.docx')

email="""Subject: Reviewer-round reanalysis and revision package

Dear Juliane,

Thank you for sending the revised metadata table and for confirming the co-last-author arrangement.

I have now completed the reviewer-requested reanalysis and prepared a revision package. I followed the new Excel table as the authoritative metadata source, retained all 73 samples from 72 patients, and did not change any disease-group assignment. The two samples from patient B2018.30786 were kept together in the patient-aware machine-learning analysis.

Please begin with 01_Juliane_revision_report.html. It links each reviewer concern to our agreed direction, the completed analysis, its plain-language interpretation, and the corresponding manuscript action. The HTML file is self-contained and can be opened offline in a browser.

The main conclusion is that the methylation data contain strong disease-group-associated structure, but source, Sentrix, demographic variables and biopsy site are also associated with the leading structure and cannot be fully separated from disease group in this cohort. The exact t-SNE geometry is parameter-dependent, and ALS versus non-ALS NMA is the least stable comparison. I therefore recommend retaining the work as a pilot, hypothesis-generating study while removing disease-entity and clinical-diagnostic overstatements.

The package contains:
- the reviewer-to-solution report;
- a clean revised manuscript draft with unresolved clinical items highlighted;
- a point-by-point rebuttal draft;
- the analysis workbook;
- the proposed figure map; and
- PDF versions of the main candidate figures for review.

I still need your decision or clinical information on the following points:
1. Archive candidate totals, exclusion counts and reasons, the exact screening flow, and replacement text for the remaining clinical-definition placeholders (`xxx`) in the Human samples section.
2. Whether all candidates in the eight-year interval were reviewed, including slide review, clinical-information review and pathologist adjudication.
3. The definition and provenance of the low/medium/high lymphomonocyte categories.
4. Whether validated cell fractions or fibrosis, necrosis, denervation, fiber-type or pathology-severity measures are available. If not, I recommend stating explicitly that cell-composition-adjusted models were unavailable and that this is a major limitation.
5. Approval of the proposed Figure 1 revision: retain the baseline t-SNE as panel E, use the label-free full-matrix correlation heatmap as panel F, and move the label-informed top-CpG heatmap to the supplement as a supervised descriptive figure.
6. Confirmation of the final author list and order, including the current Tayfun Palaz placeholder, and the exact equal-contribution/co-last-author wording.

Once these points are resolved, I can integrate your wording into the manuscript and rebuttal and prepare the submission-ready files.

Best,
Joseph
"""
(OUT/'Email_to_Juliane_draft.txt').write_text(email)
edoc=Document(); edoc.add_heading('Email to Juliane — draft',0)
for para in email.split('\n\n'): edoc.add_paragraph(para)
edoc.save(OUT/'Email_to_Juliane_draft.docx')
print(OUT)
